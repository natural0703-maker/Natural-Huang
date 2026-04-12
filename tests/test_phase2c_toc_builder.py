from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, convert
from src.phase2_toc_builder import (
    TOC_STATUS_FALLBACK_CHAPTER_LIST,
    TOC_STATUS_FIELD_INSERTED,
    TOC_STATUS_NOT_REQUESTED,
    TOC_STATUS_SKIPPED_EXISTING_TOC,
    TOC_STATUS_SKIPPED_NO_HEADINGS,
    insert_minimal_toc,
)
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[tuple[str, str | None]]) -> None:
    document = Document()
    for text, style_name in paragraphs:
        paragraph = document.add_paragraph(text)
        if style_name is not None:
            paragraph.style = style_name
    document.save(path)


def _read_docx_text(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def _write_phase2c_config(tmp: Path) -> Path:
    low_path = tmp / "low.yaml"
    high_path = tmp / "high.yaml"
    config_path = tmp / "config.yaml"
    low_path.write_text("[]\n", encoding="utf-8")
    high_path.write_text(
        '- term: "zzzz_never_match"\n'
        "  risk_category: wording\n"
        "  enabled: false\n",
        encoding="utf-8",
    )
    config_path.write_text(
        "opencc_config: s2t\n"
        "default_report_name: report.xlsx\n"
        f"default_term_dict_path: {low_path.as_posix()}\n"
        f"default_high_risk_rules_path: {high_path.as_posix()}\n"
        "enable_space_cleanup: true\n"
        "document_format:\n"
        "  page_margin_top_cm: 1.27\n"
        "  page_margin_bottom_cm: 1.27\n"
        "  page_margin_left_cm: 1.27\n"
        "  page_margin_right_cm: 1.27\n"
        "  body_first_line_indent_chars: 2\n"
        "  heading_style_name: Heading 2\n",
        encoding="utf-8",
    )
    return config_path


def test_insert_minimal_toc_inserts_word_field_for_heading_2() -> None:
    document = Document()
    heading = document.add_paragraph("Chapter 1")
    heading.style = "Heading 2"
    document.add_paragraph("Body")

    toc = insert_minimal_toc(document, requested=True)

    assert toc.status == TOC_STATUS_FIELD_INSERTED
    assert toc.fallback_used is False
    assert toc.chapter_count == 1
    assert [paragraph.text for paragraph in document.paragraphs][:3] == ["目錄", "請在 Word 中更新目錄。", "Chapter 1"]
    assert "TOC" in document.paragraphs[1]._p.xml


def test_insert_minimal_toc_skips_without_heading_2() -> None:
    document = Document()
    document.add_paragraph("Chapter 1")
    before = [paragraph.text for paragraph in document.paragraphs]

    toc = insert_minimal_toc(document, requested=True)

    assert toc.status == TOC_STATUS_SKIPPED_NO_HEADINGS
    assert toc.chapter_count == 0
    assert [paragraph.text for paragraph in document.paragraphs] == before


def test_insert_minimal_toc_skips_existing_toc_title() -> None:
    document = Document()
    document.add_paragraph("目錄")
    heading = document.add_paragraph("Chapter 1")
    heading.style = "Heading 2"

    toc = insert_minimal_toc(document, requested=True)

    assert toc.status == TOC_STATUS_SKIPPED_EXISTING_TOC
    assert [paragraph.text for paragraph in document.paragraphs].count("目錄") == 1


def test_insert_minimal_toc_falls_back_to_plain_chapter_list(monkeypatch) -> None:
    document = Document()
    heading = document.add_paragraph("Chapter 1")
    heading.style = "Heading 2"
    document.add_paragraph("Body")

    def fail_field_insert(document):
        raise RuntimeError("field failed")

    monkeypatch.setattr("src.phase2_toc_builder._insert_toc_field", fail_field_insert)

    toc = insert_minimal_toc(document, requested=True)

    assert toc.status == TOC_STATUS_FALLBACK_CHAPTER_LIST
    assert toc.fallback_used is True
    assert [paragraph.text for paragraph in document.paragraphs][:3] == ["目錄", "Chapter 1", "Chapter 1"]


def test_insert_minimal_toc_not_requested_does_nothing() -> None:
    document = Document()
    heading = document.add_paragraph("Chapter 1")
    heading.style = "Heading 2"

    toc = insert_minimal_toc(document, requested=False)

    assert toc.status == TOC_STATUS_NOT_REQUESTED
    assert [paragraph.text for paragraph in document.paragraphs] == ["Chapter 1"]


def test_convert_can_insert_toc_without_changing_body_relative_order() -> None:
    tmp = make_test_dir("phase2c_convert_toc")
    input_path = tmp / "source.docx"
    output_dir = tmp / "out"
    config_path = _write_phase2c_config(tmp)
    _make_docx(input_path, [("Chapter 1", "Heading 2"), ("Body 1", None), ("Chapter 2", "Heading 2")])

    result = convert(
        Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path, create_toc=True)
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_FIELD_INSERTED
    assert result.schema.toc.chapter_count == 2
    assert result.output_path is not None
    assert _read_docx_text(result.output_path)[:5] == ["目錄", "請在 Word 中更新目錄。", "Chapter 1", "Body 1", "Chapter 2"]


def test_convert_respects_no_create_toc() -> None:
    tmp = make_test_dir("phase2c_convert_no_toc")
    input_path = tmp / "source.docx"
    output_dir = tmp / "out"
    config_path = _write_phase2c_config(tmp)
    _make_docx(input_path, [("Chapter 1", "Heading 2"), ("Body 1", None)])

    result = convert(
        Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path, create_toc=False)
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_NOT_REQUESTED
    assert result.output_path is not None
    assert _read_docx_text(result.output_path) == ["Chapter 1", "Body 1"]
