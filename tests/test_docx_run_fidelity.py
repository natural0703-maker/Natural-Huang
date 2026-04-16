from __future__ import annotations

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

from src.docx_run_fidelity import (
    RUN_REPLACE_APPLIED,
    RUN_REPLACE_EMPTY_REPLACEMENT,
    RUN_REPLACE_EXPECTED_TEXT_MISMATCH,
    RUN_REPLACE_INVALID_SPAN,
    RUN_REPLACE_SPANS_MULTIPLE_RUNS,
    RUN_REPLACE_UNSAFE_LINE_BREAK,
    RunReplaceResult,
    try_replace_text_in_single_run,
)


def _paragraph_with_single_run(text: str):
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    return paragraph, run


def test_single_run_plain_text_replacement_succeeds() -> None:
    paragraph, _run = _paragraph_with_single_run("abc target xyz")
    run_count = len(paragraph.runs)

    result = try_replace_text_in_single_run(
        paragraph,
        4,
        10,
        "done",
        expected_text="target",
    )

    assert result == RunReplaceResult(
        status="applied",
        code=RUN_REPLACE_APPLIED,
        matched_text="target",
        replacement_text="done",
    )
    assert paragraph.text == "abc done xyz"
    assert len(paragraph.runs) == run_count


def test_single_run_replacement_preserves_common_inline_formatting() -> None:
    paragraph, run = _paragraph_with_single_run("target")
    run.bold = True
    run.italic = True
    run.underline = True
    run.font.size = Pt(14)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    result = try_replace_text_in_single_run(paragraph, 0, 6, "done")

    assert result.status == "applied"
    assert result.code == RUN_REPLACE_APPLIED
    assert paragraph.text == "done"
    assert len(paragraph.runs) == 1
    assert run.bold is True
    assert run.italic is True
    assert run.underline is True
    assert run.font.size == Pt(14)
    assert run.font.name == "Arial"
    assert run.font.color.rgb == RGBColor(0x12, 0x34, 0x56)
    assert run.font.highlight_color == WD_COLOR_INDEX.YELLOW


def test_multi_run_replacement_changes_only_target_run_and_keeps_run_count() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("before ")
    target = paragraph.add_run("target")
    after = paragraph.add_run(" after")
    first.bold = True
    target.italic = True
    after.underline = True
    run_count = len(paragraph.runs)

    result = try_replace_text_in_single_run(paragraph, 7, 13, "done")

    assert result.status == "applied"
    assert paragraph.text == "before done after"
    assert len(paragraph.runs) == run_count
    assert first.text == "before "
    assert first.bold is True
    assert target.text == "done"
    assert target.italic is True
    assert after.text == " after"
    assert after.underline is True


def test_span_crossing_runs_is_not_modified() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("abc ")
    paragraph.add_run("target")
    original_runs = [run.text for run in paragraph.runs]

    result = try_replace_text_in_single_run(paragraph, 2, 8, "done")

    assert result.status == "unsafe"
    assert result.code == RUN_REPLACE_SPANS_MULTIPLE_RUNS
    assert result.matched_text == "c targ"
    assert result.replacement_text == "done"
    assert [run.text for run in paragraph.runs] == original_runs


def test_invalid_span_is_not_modified() -> None:
    paragraph, _run = _paragraph_with_single_run("target")

    result = try_replace_text_in_single_run(paragraph, 3, 3, "done")

    assert result.status == "invalid"
    assert result.code == RUN_REPLACE_INVALID_SPAN
    assert paragraph.text == "target"


def test_empty_replacement_is_not_modified() -> None:
    paragraph, _run = _paragraph_with_single_run("target")

    result = try_replace_text_in_single_run(paragraph, 0, 6, "")

    assert result.status == "invalid"
    assert result.code == RUN_REPLACE_EMPTY_REPLACEMENT
    assert paragraph.text == "target"


def test_expected_text_mismatch_is_not_modified() -> None:
    paragraph, _run = _paragraph_with_single_run("target")

    result = try_replace_text_in_single_run(
        paragraph,
        0,
        6,
        "done",
        expected_text="source",
    )

    assert result.status == "skipped"
    assert result.code == RUN_REPLACE_EXPECTED_TEXT_MISMATCH
    assert result.matched_text == "target"
    assert result.replacement_text == "done"
    assert paragraph.text == "target"


def test_run_with_line_break_is_unsafe_and_not_modified() -> None:
    paragraph, _run = _paragraph_with_single_run("alpha\nbeta")
    original_text = paragraph.text
    original_runs = [run.text for run in paragraph.runs]

    result = try_replace_text_in_single_run(
        paragraph,
        0,
        5,
        "done",
        expected_text="alpha",
    )

    assert result.status == "unsafe"
    assert result.code == RUN_REPLACE_UNSAFE_LINE_BREAK
    assert result.matched_text == "alpha"
    assert paragraph.text == original_text
    assert [run.text for run in paragraph.runs] == original_runs


@pytest.mark.parametrize(
    ("char_start", "char_end", "replacement_text", "expected_text"),
    [
        (2, 8, "full", "target"),
        (2, 5, "start", "tar"),
        (5, 8, "end", "get"),
    ],
)
def test_multi_run_boundary_spans_inside_one_run_succeed(
    char_start: int,
    char_end: int,
    replacement_text: str,
    expected_text: str,
) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    before = paragraph.add_run("aa")
    target = paragraph.add_run("target")
    after = paragraph.add_run("zz")
    before.bold = True
    target.italic = True
    after.underline = True
    run_count = len(paragraph.runs)

    result = try_replace_text_in_single_run(
        paragraph,
        char_start,
        char_end,
        replacement_text,
        expected_text=expected_text,
    )

    assert result.status == "applied"
    assert result.code == RUN_REPLACE_APPLIED
    assert result.matched_text == expected_text
    assert len(paragraph.runs) == run_count
    assert before.text == "aa"
    assert before.bold is True
    assert target.italic is True
    assert after.text == "zz"
    assert after.underline is True


def test_span_one_character_across_run_boundary_is_not_modified() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("abc")
    second = paragraph.add_run("def")
    original_runs = [run.text for run in paragraph.runs]

    result = try_replace_text_in_single_run(paragraph, 2, 4, "X")

    assert result.status == "unsafe"
    assert result.code == RUN_REPLACE_SPANS_MULTIPLE_RUNS
    assert result.matched_text == "cd"
    assert first.text == "abc"
    assert second.text == "def"
    assert [run.text for run in paragraph.runs] == original_runs


def test_whitespace_punctuation_and_mixed_text_are_not_normalized() -> None:
    paragraph, _run = _paragraph_with_single_run("前 ABC123， target ！後")
    start = paragraph.text.index("target")
    end = start + len("target")

    result = try_replace_text_in_single_run(
        paragraph,
        start,
        end,
        "新value7",
        expected_text="target",
    )

    assert result.status == "applied"
    assert paragraph.text == "前 ABC123， 新value7 ！後"


def test_paragraph_level_style_is_not_overwritten() -> None:
    document = Document()
    style = document.styles.add_style("RunFidelityParagraphStyle", WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.add_run("before target after")

    result = try_replace_text_in_single_run(paragraph, 7, 13, "done")

    assert result.status == "applied"
    assert paragraph.text == "before done after"
    assert paragraph.style.name == "RunFidelityParagraphStyle"


def test_common_formatting_combinations_are_preserved() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    bold_color = paragraph.add_run("boldcolor")
    bold_color.bold = True
    bold_color.font.color.rgb = RGBColor(0xAA, 0x22, 0x22)
    paragraph.add_run(" ")
    italic_underline = paragraph.add_run("italicunderline")
    italic_underline.italic = True
    italic_underline.underline = True
    paragraph.add_run(" ")
    size_face = paragraph.add_run("sizeface")
    size_face.font.size = Pt(16)
    size_face.font.name = "Arial"
    paragraph.add_run(" ")
    highlight_underline = paragraph.add_run("highlightunderline")
    highlight_underline.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    highlight_underline.underline = True

    first_result = try_replace_text_in_single_run(paragraph, 0, 9, "one")
    second_start = paragraph.text.index("italicunderline")
    second_result = try_replace_text_in_single_run(
        paragraph,
        second_start,
        second_start + len("italicunderline"),
        "two",
    )
    third_start = paragraph.text.index("sizeface")
    third_result = try_replace_text_in_single_run(
        paragraph,
        third_start,
        third_start + len("sizeface"),
        "three",
    )
    fourth_start = paragraph.text.index("highlightunderline")
    fourth_result = try_replace_text_in_single_run(
        paragraph,
        fourth_start,
        fourth_start + len("highlightunderline"),
        "four",
    )

    assert first_result.status == "applied"
    assert second_result.status == "applied"
    assert third_result.status == "applied"
    assert fourth_result.status == "applied"
    assert bold_color.text == "one"
    assert bold_color.bold is True
    assert bold_color.font.color.rgb == RGBColor(0xAA, 0x22, 0x22)
    assert italic_underline.text == "two"
    assert italic_underline.italic is True
    assert italic_underline.underline is True
    assert size_face.text == "three"
    assert size_face.font.size == Pt(16)
    assert size_face.font.name == "Arial"
    assert highlight_underline.text == "four"
    assert highlight_underline.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN
    assert highlight_underline.underline is True
