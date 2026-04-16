from docx import Document
from docx.oxml import OxmlElement

from src.docx_line_break_cleanup import (
    LINE_BREAK_CLEANUP_ELIGIBLE,
    LINE_BREAK_CLEANUP_SKIPPED_DIALOGUE_LIKE,
    LINE_BREAK_CLEANUP_SKIPPED_EMPTY,
    LINE_BREAK_CLEANUP_SKIPPED_HEADING,
    LINE_BREAK_CLEANUP_SKIPPED_LIST_LIKE,
    LINE_BREAK_CLEANUP_SKIPPED_NO_LINE_BREAK,
    LINE_BREAK_CLEANUP_SKIPPED_NON_CJK_DOMINANT,
    LINE_BREAK_CLEANUP_SKIPPED_SPECIAL_STRUCTURE,
    LINE_BREAK_CLEANUP_SKIPPED_TABLE_CELL,
    apply_line_break_cleanup_to_paragraph,
    assess_paragraph_for_line_break_cleanup,
)


def test_ordinary_cjk_paragraph_with_hard_line_break_is_eligible() -> None:
    document = Document()
    paragraph = document.add_paragraph("曾經在大學的時候，王琪是劉斌心中的女神，可現\n在仍然如此。")

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is True
    assert result.code == LINE_BREAK_CLEANUP_ELIGIBLE


def test_paragraph_without_line_break_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("這是一個普通段落，不需要清理。")

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_NO_LINE_BREAK


def test_empty_paragraph_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("  \n  ")

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_EMPTY


def test_heading_paragraph_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("第一章\n重逢")
    paragraph.style = "Heading 2"

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_HEADING


def test_table_cell_paragraph_is_skipped() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.text = "表格內文字\n不應清理"

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_TABLE_CELL


def test_text_list_like_paragraph_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("1. 第一項\n2. 第二項")

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_LIST_LIKE


def test_word_list_style_paragraph_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("第一項\n第二項")
    paragraph.style = "List Bullet"

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_LIST_LIKE


def test_dialogue_like_paragraph_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("「你還好嗎？」\n「我沒事。」")

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_DIALOGUE_LIKE


def test_english_dominant_paragraph_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("This paragraph has a hard\nline break but is English.")

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_NON_CJK_DOMINANT


def test_hyperlink_like_structure_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("這段文字\n含連結")
    hyperlink = OxmlElement("w:hyperlink")
    run = paragraph.runs[0]._r
    paragraph._p.remove(run)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_SPECIAL_STRUCTURE


def test_field_like_structure_is_skipped() -> None:
    document = Document()
    paragraph = document.add_paragraph("這段文字\n含欄位")
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:fldChar"))

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is False
    assert result.code == LINE_BREAK_CLEANUP_SKIPPED_SPECIAL_STRUCTURE


def test_selector_does_not_modify_paragraph_text() -> None:
    document = Document()
    paragraph = document.add_paragraph("王琪是劉斌心中的女神，可現\n在仍然如此。")
    original_text = paragraph.text

    result = assess_paragraph_for_line_break_cleanup(paragraph)

    assert result.eligible is True
    assert paragraph.text == original_text


def test_cleanup_wrapper_applies_to_eligible_paragraph() -> None:
    document = Document()
    paragraph = document.add_paragraph("王琪是劉斌心中的女神，可現\n在仍然如此。")

    result = apply_line_break_cleanup_to_paragraph(paragraph)

    assert result.changed is True
    assert result.eligibility_code == LINE_BREAK_CLEANUP_ELIGIBLE
    assert result.cleanup_applied_count == 1
    assert result.cleanup_skipped_count == 0
    assert paragraph.text == "王琪是劉斌心中的女神，可現在仍然如此。"


def test_cleanup_wrapper_does_not_modify_ineligible_paragraph() -> None:
    document = Document()
    paragraph = document.add_paragraph("第一章\n重逢")
    paragraph.style = "Heading 2"
    original_text = paragraph.text

    result = apply_line_break_cleanup_to_paragraph(paragraph)

    assert result.changed is False
    assert result.eligibility_code == LINE_BREAK_CLEANUP_SKIPPED_HEADING
    assert result.cleanup_applied_count == 0
    assert result.cleanup_skipped_count == 0
    assert paragraph.text == original_text


def test_cleanup_wrapper_keeps_text_when_plain_text_helper_skips_boundary() -> None:
    document = Document()
    paragraph = document.add_paragraph("這是一個完整句子。\n下一句也應保留換行。")
    original_text = paragraph.text

    result = apply_line_break_cleanup_to_paragraph(paragraph)

    assert result.changed is False
    assert result.eligibility_code == LINE_BREAK_CLEANUP_ELIGIBLE
    assert result.cleanup_applied_count == 0
    assert result.cleanup_skipped_count == 1
    assert paragraph.text == original_text
