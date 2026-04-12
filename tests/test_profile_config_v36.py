from pathlib import Path

from docx import Document

from src.cli_v35 import main as cli_main
from src.config_loader import load_config
from tests.test_paths import make_test_dir


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _write_profile_rule_files(tmp: Path) -> None:
    (tmp / "low_default.yaml").write_text(
        "- source: 支持\n"
        "  target: 支援\n"
        "  risk_level: low\n"
        "  enabled: true\n"
        "  category: wording\n"
        "  note: default\n",
        encoding="utf-8",
    )
    (tmp / "low_strict.yaml").write_text(
        "- source: 支持\n"
        "  target: 支援\n"
        "  risk_level: low\n"
        "  enabled: true\n"
        "  category: wording\n"
        "  note: strict\n",
        encoding="utf-8",
    )
    (tmp / "high_default.yaml").write_text(
        "- term: 資訊\n"
        "  risk_category: wording\n"
        "  enabled: true\n"
        "  suggested_candidates: 請人工確認\n",
        encoding="utf-8",
    )
    (tmp / "high_strict.yaml").write_text(
        "- term: 支持\n"
        "  risk_category: wording\n"
        "  enabled: true\n"
        "  suggested_candidates: 請人工確認\n",
        encoding="utf-8",
    )


def _write_profile_config(tmp: Path) -> Path:
    config = tmp / "config.yaml"
    config.write_text(
        "opencc_config: s2t\n"
        "default_report_name: report.xlsx\n"
        "enable_space_cleanup: true\n"
        "active_profile: default\n"
        "profiles:\n"
        "  default:\n"
        "    description: 預設方案\n"
        "    low_risk_dict: low_default.yaml\n"
        "    high_risk_rules: high_default.yaml\n"
        "    format_config:\n"
        "      body_font_size_pt: 12\n"
        "  strict_tw:\n"
        "    description: 嚴格方案\n"
        "    low_risk_dict: low_strict.yaml\n"
        "    high_risk_rules: high_strict.yaml\n"
        "    format_config:\n"
        "      body_font_size_pt: 14\n",
        encoding="utf-8",
    )
    return config


def test_old_config_still_works_without_profiles() -> None:
    tmp = make_test_dir("v36_old_config_compat")
    config = tmp / "config.yaml"
    term_dict = tmp / "term_dict.yaml"
    high_rules = tmp / "high.yaml"
    term_dict.write_text('"支持": "支援"\n', encoding="utf-8")
    high_rules.write_text(
        "- term: 資訊\n"
        "  risk_category: wording\n"
        "  enabled: true\n"
        "  suggested_candidates: 請人工確認\n",
        encoding="utf-8",
    )
    config.write_text(
        "opencc_config: s2t\n"
        "default_report_name: report.xlsx\n"
        "default_term_dict_path: term_dict.yaml\n"
        "default_high_risk_rules_path: high.yaml\n"
        "enable_space_cleanup: true\n",
        encoding="utf-8",
    )

    loaded = load_config(config)
    assert loaded.active_profile == "default"
    assert loaded.default_term_dict_path.name == "term_dict.yaml"
    assert loaded.default_high_risk_rules_path.name == "high.yaml"


def test_active_profile_loads_correct_paths() -> None:
    tmp = make_test_dir("v36_active_profile")
    _write_profile_rule_files(tmp)
    config = _write_profile_config(tmp)

    loaded = load_config(config)
    assert loaded.active_profile == "default"
    assert loaded.default_term_dict_path.name == "low_default.yaml"
    assert loaded.default_high_risk_rules_path.name == "high_default.yaml"
    assert sorted(loaded.profiles.keys()) == ["default", "strict_tw"]


def test_profile_override_argument_changes_selection() -> None:
    tmp = make_test_dir("v36_profile_override")
    _write_profile_rule_files(tmp)
    config = _write_profile_config(tmp)

    loaded = load_config(config, profile_name="strict_tw")
    assert loaded.active_profile == "strict_tw"
    assert loaded.default_term_dict_path.name == "low_strict.yaml"
    assert loaded.default_high_risk_rules_path.name == "high_strict.yaml"
    assert loaded.document_format.body_font_size_pt == 14.0


def test_profile_not_found_raises_clear_error() -> None:
    tmp = make_test_dir("v36_profile_not_found")
    _write_profile_rule_files(tmp)
    config = _write_profile_config(tmp)

    try:
        load_config(config, profile_name="missing_profile")
    except ValueError as exc:
        assert "找不到 profile" in str(exc)
        return
    raise AssertionError("預期不存在的 profile 應拋出 ValueError")


def test_profile_target_file_missing_raises_error() -> None:
    tmp = make_test_dir("v36_profile_missing_file")
    _write_profile_rule_files(tmp)
    config = tmp / "config.yaml"
    config.write_text(
        "opencc_config: s2t\n"
        "default_report_name: report.xlsx\n"
        "enable_space_cleanup: true\n"
        "active_profile: broken\n"
        "profiles:\n"
        "  broken:\n"
        "    low_risk_dict: low_default.yaml\n"
        "    high_risk_rules: missing_high.yaml\n",
        encoding="utf-8",
    )

    try:
        load_config(config)
    except FileNotFoundError as exc:
        assert "找不到 profile `broken` 的高風險規則檔案" in str(exc)
        return
    raise AssertionError("預期缺少規則檔時應拋出 FileNotFoundError")


def test_cli_profile_switch_changes_behavior() -> None:
    tmp = make_test_dir("v36_cli_profile_switch")
    _write_profile_rule_files(tmp)
    config = _write_profile_config(tmp)
    input_docx = tmp / "sample.docx"
    _write_docx(input_docx, ["我們支持這個方案。"])

    default_out = tmp / "out_default"
    strict_out = tmp / "out_strict"
    default_out.mkdir(parents=True, exist_ok=True)
    strict_out.mkdir(parents=True, exist_ok=True)

    exit_default = cli_main(
        [
            "--input-file",
            str(input_docx),
            "--output-dir",
            str(default_out),
            "--config",
            str(config),
            "--profile",
            "default",
        ]
    )
    exit_strict = cli_main(
        [
            "--input-file",
            str(input_docx),
            "--output-dir",
            str(strict_out),
            "--config",
            str(config),
            "--profile",
            "strict_tw",
        ]
    )
    assert exit_default == 0
    assert exit_strict == 0

    default_doc = Document(str(default_out / "sample_TW.docx"))
    strict_doc = Document(str(strict_out / "sample_TW.docx"))
    default_text = "\n".join(p.text for p in default_doc.paragraphs)
    strict_text = "\n".join(p.text for p in strict_doc.paragraphs)
    assert "支援" in default_text
    assert "支持" in strict_text


def test_cli_profile_not_found_returns_error() -> None:
    tmp = make_test_dir("v36_cli_missing_profile")
    _write_profile_rule_files(tmp)
    config = _write_profile_config(tmp)
    input_docx = tmp / "sample.docx"
    _write_docx(input_docx, ["我們支持這個方案。"])
    output_dir = tmp / "out"
    output_dir.mkdir(parents=True, exist_ok=True)

    exit_code = cli_main(
        [
            "--input-file",
            str(input_docx),
            "--output-dir",
            str(output_dir),
            "--config",
            str(config),
            "--profile",
            "not_exists",
        ]
    )
    assert exit_code == 1
