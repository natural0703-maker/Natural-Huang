from pathlib import Path

from docx import Document

from src.docx_line_break_cleanup import (
    LINE_BREAK_CLEANUP_ELIGIBLE,
    LINE_BREAK_CLEANUP_SKIPPED_HEADING,
    LINE_BREAK_CLEANUP_SKIPPED_LIST_LIKE,
    LINE_BREAK_CLEANUP_SKIPPED_NO_LINE_BREAK,
    LINE_BREAK_CLEANUP_SKIPPED_NON_CJK_DOMINANT,
)
from src.phase1_config import load_phase1_config
from src.phase1_converter import convert_docx
from src.phase1_pipeline import Phase1Options, convert


def _make_config(tmp_path: Path) -> Path:
    low_path = tmp_path / "low.yaml"
    high_path = tmp_path / "high.yaml"
    config_path = tmp_path / "config.yaml"
    low_path.write_text("{}\n", encoding="utf-8")
    high_path.write_text(
        "- term: __never_match_line_break_cleanup__\n"
        "  risk_category: wording\n"
        "  enabled: false\n",
        encoding="utf-8",
    )
    config_path.write_text(
        "opencc_config: s2t\n"
        "default_report_name: report.xlsx\n"
        f"default_term_dict_path: {low_path.as_posix()}\n"
        f"default_high_risk_rules_path: {high_path.as_posix()}\n"
        "enable_space_cleanup: true\n"
        "document_format:\n"
        "  page_margin_top_cm: 1.27\n"
        "  page_margin_bottom_cm: 1.27\n"
        "  page_margin_left_cm: 1.27\n"
        "  page_margin_right_cm: 1.27\n"
        "  body_first_line_indent_chars: 2\n"
        "  heading_style_name: Heading 2\n",
        encoding="utf-8",
    )
    return config_path


def _save_doc(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _read_paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def test_convert_line_break_cleanup_is_disabled_by_default(tmp_path: Path) -> None:
    tmp = tmp_path / "line_break_convert_default"
    tmp.mkdir()
    input_path = tmp / "input.docx"
    output_dir = tmp / "out"
    config_path = _make_config(tmp)
    source = "曾經在大學的時候，王琪是劉斌心中的女神，成績雖算不上十分優秀，可現\n在有哪個男生找女友是在意女生的成績好壞？"
    _save_doc(input_path, [source])

    result = convert(Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path, create_toc=False))

    assert result.docx_processed is True
    assert result.output_path is not None
    assert _read_paragraphs(result.output_path) == [source]


def test_convert_line_break_cleanup_opt_in_cleans_ordinary_paragraph(tmp_path: Path) -> None:
    tmp = tmp_path / "line_break_convert_enabled"
    tmp.mkdir()
    input_path = tmp / "input.docx"
    output_dir = tmp / "out"
    config_path = _make_config(tmp)
    source = "曾經在大學的時候，王琪是劉斌心中的女神，成績雖算不上十分優秀，可現\n在有哪個男生找女友是在意女生的成績好壞？"
    expected = "曾經在大學的時候，王琪是劉斌心中的女神，成績雖算不上十分優秀，可現在有哪個男生找女友是在意女生的成績好壞？"
    _save_doc(input_path, [source])

    result = convert(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            config_path=config_path,
            create_toc=False,
            enable_line_break_cleanup=True,
        )
    )

    assert result.docx_processed is True
    assert result.output_path is not None
    assert _read_paragraphs(result.output_path) == [expected]


def test_convert_line_break_cleanup_summary_is_available_from_converter(tmp_path: Path) -> None:
    tmp = tmp_path / "line_break_convert_summary"
    tmp.mkdir()
    input_path = tmp / "input.docx"
    output_dir = tmp / "out"
    config_path = _make_config(tmp)
    document = Document()
    document.add_paragraph("曾經在大學的時候，王琪是劉斌心中的女神，成績雖算不上十分優秀，可現\n在有哪個男生找女友是在意女生的成績好壞？")
    heading = document.add_paragraph("第一章\n開始")
    heading.style = "Heading 2"
    document.add_paragraph("1. 第一項\n2. 第二項")
    document.add_paragraph("This paragraph has a hard\nline break.")
    document.add_paragraph("這是一個沒有硬斷行的普通段落。")
    document.save(input_path)

    result = convert_docx(
        input_path,
        output_dir,
        load_phase1_config(config_path),
        create_toc=False,
        enable_line_break_cleanup=True,
    )

    summary = result.line_break_cleanup_summary
    assert result.converted is True
    assert summary.scanned_count == 5
    assert summary.eligible_count == 1
    assert summary.changed_count == 1
    assert summary.skipped_count == 4
    assert summary.codes[LINE_BREAK_CLEANUP_ELIGIBLE] == 1
    assert summary.codes[LINE_BREAK_CLEANUP_SKIPPED_HEADING] == 1
    assert summary.codes[LINE_BREAK_CLEANUP_SKIPPED_LIST_LIKE] == 1
    assert summary.codes[LINE_BREAK_CLEANUP_SKIPPED_NON_CJK_DOMINANT] == 1
    assert summary.codes[LINE_BREAK_CLEANUP_SKIPPED_NO_LINE_BREAK] == 1


def test_convert_line_break_cleanup_keeps_high_risk_paragraphs_and_tables_unchanged(tmp_path: Path) -> None:
    tmp = tmp_path / "line_break_convert_skips"
    tmp.mkdir()
    input_path = tmp / "input.docx"
    output_dir = tmp / "out"
    config_path = _make_config(tmp)
    document = Document()
    heading = document.add_paragraph("第一章\n開始")
    heading.style = "Heading 2"
    document.add_paragraph("1. 第一項\n2. 第二項")
    document.add_paragraph("This paragraph has a hard\nline break.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "表格中的中文硬\n斷行不應被處理"
    document.save(input_path)

    result = convert(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            config_path=config_path,
            create_toc=False,
            enable_line_break_cleanup=True,
        )
    )

    assert result.output_path is not None
    output_doc = Document(result.output_path)
    assert [paragraph.text for paragraph in output_doc.paragraphs] == [
        "第一章\n開始",
        "1. 第一項\n2. 第二項",
        "This paragraph has a hard\nline break.",
    ]
    assert output_doc.tables[0].cell(0, 0).paragraphs[0].text == "表格中的中文硬\n斷行不應被處理"


def test_convert_line_break_cleanup_output_docx_can_be_reopened(tmp_path: Path) -> None:
    tmp = tmp_path / "line_break_convert_reopen"
    tmp.mkdir()
    input_path = tmp / "input.docx"
    output_dir = tmp / "out"
    config_path = _make_config(tmp)
    _save_doc(input_path, ["曾經在大學的時候，王琪是劉斌心中的女神，成績雖算不上十分優秀，可現\n在有哪個男生找女友是在意女生的成績好壞？"])

    result = convert(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            config_path=config_path,
            create_toc=False,
            enable_line_break_cleanup=True,
        )
    )

    assert result.output_path is not None
    reopened = Document(result.output_path)
    assert reopened.paragraphs
