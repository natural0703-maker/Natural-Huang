from docx import Document

from src.config_loader import DocumentFormatConfig
from src.docx_writer import ParagraphOutput, write_paragraphs_to_docx
from tests.test_paths import make_test_dir


def test_writer_applies_heading_and_body_format() -> None:
    tmp = make_test_dir("docx_format_v31")
    out_path = tmp / "out.docx"

    fmt = DocumentFormatConfig()
    write_paragraphs_to_docx(
        [
            ParagraphOutput("\u7b2c\u4e00\u7ae0 \u6211\u898b\u7336\u6190", is_heading=True),
            ParagraphOutput("\u9019\u662f\u4e00\u6bb5\u4e00\u822c\u5167\u6587\u3002", is_heading=False),
        ],
        out_path,
        document_format=fmt,
    )

    doc = Document(str(out_path))
    assert len(doc.paragraphs) == 2

    heading_paragraph = doc.paragraphs[0]
    body_paragraph = doc.paragraphs[1]

    assert heading_paragraph.style.name in {"Heading 2", "\u6a19\u984c 2"}
    assert body_paragraph.style.name in {"Normal", "\u6a19\u6e96", "\u5167\u6587", "\u672c\u6587"}

    body_run = body_paragraph.runs[0]
    assert body_run.font.name == fmt.body_font_name
    assert body_run.font.size is not None
    assert abs(body_run.font.size.pt - fmt.body_font_size_pt) < 0.01
