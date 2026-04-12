import json
from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, apply_review
from tests.test_paths import make_test_dir


def _make_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("alpha")
    document.add_paragraph("beta")
    document.save(path)


def _write_payload(path: Path, payload) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _candidate(**overrides):
    data = {
        "candidate_id": "paragraph_merge:0:1:stable",
        "type": "paragraph_merge",
        "status": "accepted",
        "paragraph_index": 0,
        "next_paragraph_index": 1,
        "source_text": "alpha",
        "next_source_text": "beta",
        "reason": "manual review accepted",
    }
    data.update(overrides)
    return data


def _without(data: dict, key: str) -> dict:
    return {item_key: value for item_key, value in data.items() if item_key != key}


def _apply_payload(payload):
    tmp = make_test_dir("phase2e1_paragraph_merge_schema_validation")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path)
    _write_payload(review_path, payload)
    return apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))


def _assert_flow_error(payload, code: str):
    result = _apply_payload(payload)

    assert result.docx_processed is False
    assert result.output_path is None
    assert result.schema.errors
    assert result.schema.errors[0].code == code
    assert result.apply_result is not None
    assert result.apply_result.candidate_results == []
    return result


def test_paragraph_merge_candidates_can_be_missing() -> None:
    result = _apply_payload({})

    assert result.schema.errors == []
    assert result.apply_result is not None
    assert result.apply_result.candidate_results == []
    assert result.output_path is not None
    assert result.output_path.exists()


def test_paragraph_merge_candidates_must_be_list_when_present() -> None:
    _assert_flow_error({"paragraph_merge_candidates": {}}, "PARAGRAPH_MERGE_CANDIDATES_INVALID")


def test_paragraph_merge_candidate_must_be_object() -> None:
    _assert_flow_error({"paragraph_merge_candidates": ["bad"]}, "PARAGRAPH_MERGE_CANDIDATE_INVALID")


def test_paragraph_merge_candidate_id_is_required() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_without(_candidate(), "candidate_id")]},
        "PARAGRAPH_MERGE_CANDIDATE_ID_INVALID",
    )
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(candidate_id="")]},
        "PARAGRAPH_MERGE_CANDIDATE_ID_INVALID",
    )


def test_paragraph_merge_type_must_be_string_and_supported() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(type=123)]},
        "PARAGRAPH_MERGE_CANDIDATE_TYPE_INVALID",
    )
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(type="high_risk_term")]},
        "PARAGRAPH_MERGE_CANDIDATE_TYPE_UNSUPPORTED",
    )


def test_paragraph_merge_status_must_be_supported_and_not_auto_accepted() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(status="done")]},
        "PARAGRAPH_MERGE_CANDIDATE_STATUS_INVALID",
    )
    result = _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(status="auto_accepted")]},
        "PARAGRAPH_MERGE_CANDIDATE_STATUS_INVALID",
    )

    assert "auto_accepted" in result.schema.errors[0].message


def test_paragraph_merge_paragraph_index_must_be_non_negative_integer() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(paragraph_index="0")]},
        "PARAGRAPH_MERGE_PARAGRAPH_INDEX_INVALID",
    )
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(paragraph_index=-1)]},
        "PARAGRAPH_MERGE_PARAGRAPH_INDEX_INVALID",
    )


def test_paragraph_merge_next_paragraph_index_must_be_non_negative_integer() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(next_paragraph_index="1")]},
        "PARAGRAPH_MERGE_NEXT_PARAGRAPH_INDEX_INVALID",
    )
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(next_paragraph_index=-1)]},
        "PARAGRAPH_MERGE_NEXT_PARAGRAPH_INDEX_INVALID",
    )


def test_paragraph_merge_candidates_must_be_adjacent() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(next_paragraph_index=2)]},
        "PARAGRAPH_MERGE_NOT_ADJACENT",
    )


def test_paragraph_merge_source_text_is_required() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_without(_candidate(), "source_text")]},
        "PARAGRAPH_MERGE_SOURCE_TEXT_INVALID",
    )
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(source_text="")]},
        "PARAGRAPH_MERGE_SOURCE_TEXT_INVALID",
    )


def test_paragraph_merge_next_source_text_is_required() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_without(_candidate(), "next_source_text")]},
        "PARAGRAPH_MERGE_NEXT_SOURCE_TEXT_INVALID",
    )
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(next_source_text="")]},
        "PARAGRAPH_MERGE_NEXT_SOURCE_TEXT_INVALID",
    )


def test_paragraph_merge_reason_must_be_string_when_present() -> None:
    _assert_flow_error(
        {"paragraph_merge_candidates": [_candidate(reason=["bad"])]},
        "PARAGRAPH_MERGE_REASON_INVALID",
    )


def test_paragraph_merge_supported_statuses_pass_validation() -> None:
    for status in ("accepted", "pending", "rejected", "skip"):
        result = _apply_payload({"paragraph_merge_candidates": [_candidate(status=status)]})

        assert result.schema.errors == []
        assert result.apply_result is not None
        assert result.apply_result.candidate_results == []
        assert result.output_path is not None
        assert result.output_path.exists()
