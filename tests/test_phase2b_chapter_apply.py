import json
from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, apply_review
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[str], heading2_indices: set[int] | None = None) -> None:
    document = Document()
    heading2_indices = heading2_indices or set()
    for index, text in enumerate(paragraphs):
        paragraph = document.add_paragraph(text)
        if index in heading2_indices:
            paragraph.style = "Heading 2"
    document.save(path)


def _read_docx(path: Path) -> Document:
    return Document(path)


def _write_payload(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _chapter_candidate(**overrides) -> dict:
    data = {
        "candidate_id": "chapter:0",
        "type": "chapter",
        "status": "accepted",
        "paragraph_index": 0,
    }
    data.update(overrides)
    return data


def _review_candidate(**overrides) -> dict:
    data = {
        "candidate_id": "risk:one",
        "type": "high_risk_term",
        "status": "accepted",
        "source_text": "bad",
        "resolved_text": "good",
        "paragraph_index": 0,
        "char_start": 0,
        "char_end": 3,
        "context_before": "",
        "context_after": "",
    }
    data.update(overrides)
    return data


def _apply_payload(name: str, paragraphs: list[str], payload: dict, heading2_indices: set[int] | None = None):
    tmp = make_test_dir(name)
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, paragraphs, heading2_indices=heading2_indices)
    _write_payload(review_path, payload)

    return apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))


def _result_codes(result) -> list[str]:
    assert result.apply_result is not None
    return [item.result_code for item in result.apply_result.candidate_results]


def test_accepted_chapter_candidate_applies_heading_2_without_text_or_order_changes() -> None:
    paragraphs = ["第一章 夜雨", "正文第一段"]
    result = _apply_payload(
        "phase2b_chapter_applied",
        paragraphs,
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
    )

    assert result.schema.errors == []
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert [paragraph.text for paragraph in document.paragraphs] == paragraphs
    assert document.paragraphs[0].style.name == "Heading 2"
    assert len(document.paragraphs) == len(paragraphs)
    assert "APPLIED_CHAPTER_HEADING" in _result_codes(result)


def test_pending_rejected_skip_chapter_candidates_are_not_applied() -> None:
    result = _apply_payload(
        "phase2b_chapter_status_skipped",
        ["第一章 夜雨", "第二章 晨光", "第三章 歸途"],
        {
            "chapter_candidates": [
                _chapter_candidate(candidate_id="pending", status="pending", paragraph_index=0),
                _chapter_candidate(candidate_id="rejected", status="rejected", paragraph_index=1),
                _chapter_candidate(candidate_id="skip", status="skip", paragraph_index=2),
            ]
        },
    )

    document = _read_docx(result.output_path)
    assert [paragraph.style.name for paragraph in document.paragraphs] == ["Normal", "Normal", "Normal"]
    assert _result_codes(result) == ["SKIPPED_CHAPTER_STATUS", "SKIPPED_CHAPTER_STATUS", "SKIPPED_CHAPTER_STATUS"]


def test_invalid_chapter_paragraph_index_is_candidate_level_result() -> None:
    result = _apply_payload(
        "phase2b_chapter_index_invalid",
        ["第一章 夜雨"],
        {"chapter_candidates": [_chapter_candidate(paragraph_index=5)]},
    )

    assert result.schema.errors == []
    assert _result_codes(result) == ["CHAPTER_PARAGRAPH_INDEX_INVALID"]


def test_unsupported_chapter_type_is_candidate_level_skipped() -> None:
    result = _apply_payload(
        "phase2b_chapter_unsupported_type",
        ["第一章 夜雨"],
        {"chapter_candidates": [_chapter_candidate(type="volume")]},
    )

    assert result.schema.errors == []
    assert _result_codes(result) == ["SKIPPED_UNSUPPORTED_TYPE"]


def test_chapter_apply_does_not_create_toc_or_extra_paragraphs() -> None:
    paragraphs = ["第一章 夜雨", "正文第一段"]
    result = _apply_payload(
        "phase2b_chapter_no_toc",
        paragraphs,
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
    )

    document = _read_docx(result.output_path)
    assert [paragraph.text for paragraph in document.paragraphs] == paragraphs
    assert len(document.paragraphs) == 2


def test_existing_heading_2_is_counted_as_applied() -> None:
    result = _apply_payload(
        "phase2b_chapter_already_heading",
        ["第一章 夜雨"],
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
        heading2_indices={0},
    )

    document = _read_docx(result.output_path)
    assert document.paragraphs[0].style.name == "Heading 2"
    assert _result_codes(result) == ["APPLIED_CHAPTER_HEADING"]
    assert result.apply_result.applied_count == 1


def test_duplicate_accepted_chapter_candidates_conflict_after_first_apply() -> None:
    result = _apply_payload(
        "phase2b_chapter_duplicate_conflict",
        ["第一章 夜雨"],
        {
            "chapter_candidates": [
                _chapter_candidate(candidate_id="first", paragraph_index=0),
                _chapter_candidate(candidate_id="second", paragraph_index=0),
            ]
        },
    )

    assert _result_codes(result) == ["APPLIED_CHAPTER_HEADING", "SKIPPED_CHAPTER_CONFLICT"]
    assert result.apply_result.applied_count == 1
    assert result.apply_result.skipped_count == 1


def test_missing_chapter_candidates_does_not_block_review_candidate_apply() -> None:
    result = _apply_payload(
        "phase2b_missing_chapter_candidates",
        ["bad"],
        {"review_candidates": [_review_candidate()]},
    )

    document = _read_docx(result.output_path)
    assert [paragraph.text for paragraph in document.paragraphs] == ["good"]
    assert result.schema.errors == []
    assert _result_codes(result) == ["APPLIED"]


def test_chapter_candidates_must_be_list() -> None:
    result = _apply_payload(
        "phase2b_chapter_candidates_invalid",
        ["第一章 夜雨"],
        {"chapter_candidates": {}},
    )

    assert result.docx_processed is False
    assert result.output_path is None
    assert result.schema.errors
    assert result.schema.errors[0].code == "CHAPTER_CANDIDATES_INVALID"
    assert result.apply_result is not None
    assert result.apply_result.candidate_results == []


def test_blank_chapter_paragraph_is_not_applied() -> None:
    result = _apply_payload(
        "phase2b_chapter_blank",
        [""],
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
    )

    assert _result_codes(result) == ["CHAPTER_PARAGRAPH_EMPTY"]
