from docx import Document
from docx.enum.text import WD_LINE_SPACING

from src.config_loader import DocumentFormatConfig
from src.docx_writer import ParagraphOutput, write_paragraphs_to_docx
from tests.test_paths import make_test_dir


def test_writer_applies_page_margins_and_body_paragraph_spacing() -> None:
    tmp = make_test_dir("docx_page_paragraph_v31")
    out_path = tmp / "out.docx"

    fmt = DocumentFormatConfig()
    write_paragraphs_to_docx(
        [
            ParagraphOutput("\u7b2c\u4e00\u7ae0 \u6211\u898b\u7336\u6190", is_heading=True),
            ParagraphOutput("\u9019\u662f\u4e00\u6bb5\u4e00\u822c\u5167\u6587\u3002", is_heading=False),
        ],
        out_path,
        document_format=fmt,
    )

    doc = Document(str(out_path))
    section = doc.sections[0]
    assert abs(section.top_margin.cm - fmt.page_margin_top_cm) < 0.01
    assert abs(section.bottom_margin.cm - fmt.page_margin_bottom_cm) < 0.01
    assert abs(section.left_margin.cm - fmt.page_margin_left_cm) < 0.01
    assert abs(section.right_margin.cm - fmt.page_margin_right_cm) < 0.01

    body = doc.paragraphs[1]
    pf = body.paragraph_format
    assert pf.space_before is not None
    assert pf.space_after is not None
    assert pf.line_spacing is not None
    assert abs(pf.space_before.pt - fmt.body_space_before_pt) < 0.01
    assert abs(pf.space_after.pt - fmt.body_space_after_pt) < 0.01
    assert pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST
    assert abs(pf.line_spacing.pt - fmt.body_min_line_height_pt) < 0.01
    assert pf.first_line_indent is not None
    assert abs(pf.first_line_indent.pt - (fmt.body_first_line_indent_chars * fmt.body_font_size_pt)) < 0.01

    heading = doc.paragraphs[0]
    heading_indent = heading.paragraph_format.first_line_indent
    assert heading_indent is None or abs(heading_indent.pt) < 0.01
