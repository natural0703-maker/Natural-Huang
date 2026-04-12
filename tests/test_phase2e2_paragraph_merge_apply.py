import json
from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, apply_review
from src.phase2_toc_builder import TOC_STATUS_FIELD_INSERTED
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[tuple[str, str | None]]) -> None:
    document = Document()
    for text, style_name in paragraphs:
        paragraph = document.add_paragraph(text)
        if style_name is not None:
            paragraph.style = style_name
    document.save(path)


def _read_docx(path: Path) -> Document:
    return Document(path)


def _texts(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def _write_payload(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _merge_candidate(**overrides) -> dict:
    data = {
        "candidate_id": "paragraph_merge:0:1",
        "type": "paragraph_merge",
        "status": "accepted",
        "paragraph_index": 0,
        "next_paragraph_index": 1,
        "source_text": "前段  ",
        "next_source_text": "  後段",
        "reason": "manual review accepted",
    }
    data.update(overrides)
    return data


def _apply_payload(
    name: str,
    paragraphs: list[tuple[str, str | None]],
    payload: dict,
    create_toc: bool = False,
):
    tmp = make_test_dir(name)
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, paragraphs)
    _write_payload(review_path, payload)
    return apply_review(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            apply_review_path=review_path,
            create_toc=create_toc,
        )
    )


def _result_codes(result) -> list[str]:
    assert result.apply_result is not None
    return [item.result_code for item in result.apply_result.candidate_results]


def test_accepted_merge_candidate_merges_adjacent_paragraphs() -> None:
    result = _apply_payload(
        "phase2e2_merge_applied",
        [("前段  ", None), ("  後段", None), ("第三段", None)],
        {"paragraph_merge_candidates": [_merge_candidate()]},
    )

    assert result.schema.errors == []
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert _texts(document) == ["前段後段", "第三段"]
    assert len(document.paragraphs) == 2
    assert _result_codes(result) == ["APPLIED_PARAGRAPH_MERGE"]
    assert result.apply_result.applied_count == 1


def test_pending_rejected_skip_merge_candidates_are_not_applied() -> None:
    result = _apply_payload(
        "phase2e2_merge_status_skipped",
        [("前段  ", None), ("  後段", None)],
        {
            "paragraph_merge_candidates": [
                _merge_candidate(candidate_id="pending", status="pending"),
                _merge_candidate(candidate_id="rejected", status="rejected"),
                _merge_candidate(candidate_id="skip", status="skip"),
            ]
        },
    )

    document = _read_docx(result.output_path)
    assert _texts(document) == ["前段  ", "  後段"]
    assert _result_codes(result) == [
        "SKIPPED_PARAGRAPH_MERGE_STATUS",
        "SKIPPED_PARAGRAPH_MERGE_STATUS",
        "SKIPPED_PARAGRAPH_MERGE_STATUS",
    ]
    assert result.apply_result.skipped_count == 3


def test_non_merge_type_is_rejected_by_schema_validation() -> None:
    result = _apply_payload(
        "phase2e2_merge_unsupported_type",
        [("前段  ", None), ("  後段", None)],
        {"paragraph_merge_candidates": [_merge_candidate(type="high_risk_term")]},
    )

    assert result.docx_processed is False
    assert result.output_path is None
    assert result.schema.errors
    assert result.schema.errors[0].code == "PARAGRAPH_MERGE_CANDIDATE_TYPE_UNSUPPORTED"
    assert result.apply_result.candidate_results == []


def test_invalid_index_is_candidate_level_failed() -> None:
    result = _apply_payload(
        "phase2e2_merge_invalid_index",
        [("前段  ", None), ("  後段", None)],
        {"paragraph_merge_candidates": [_merge_candidate(paragraph_index=5, next_paragraph_index=6)]},
    )

    assert result.schema.errors == []
    assert _result_codes(result) == ["PARAGRAPH_MERGE_PARAGRAPH_INDEX_INVALID"]
    assert result.apply_result.failed_count == 1


def test_non_adjacent_index_is_rejected_by_schema_validation() -> None:
    result = _apply_payload(
        "phase2e2_merge_not_adjacent",
        [("前段  ", None), ("  後段", None), ("第三段", None)],
        {"paragraph_merge_candidates": [_merge_candidate(next_paragraph_index=2)]},
    )

    assert result.docx_processed is False
    assert result.output_path is None
    assert result.schema.errors[0].code == "PARAGRAPH_MERGE_NOT_ADJACENT"


def test_blank_paragraph_is_not_merged() -> None:
    result = _apply_payload(
        "phase2e2_merge_blank_paragraph",
        [("", None), ("  後段", None)],
        {"paragraph_merge_candidates": [_merge_candidate()]},
    )

    assert _texts(_read_docx(result.output_path)) == ["", "  後段"]
    assert _result_codes(result) == ["PARAGRAPH_MERGE_PARAGRAPH_EMPTY"]


def test_source_text_mismatch_is_not_merged() -> None:
    result = _apply_payload(
        "phase2e2_merge_source_mismatch",
        [("目前前段", None), ("  後段", None)],
        {"paragraph_merge_candidates": [_merge_candidate()]},
    )

    assert _texts(_read_docx(result.output_path)) == ["目前前段", "  後段"]
    assert _result_codes(result) == ["PARAGRAPH_MERGE_SOURCE_MISMATCH"]


def test_next_source_text_mismatch_is_not_merged() -> None:
    result = _apply_payload(
        "phase2e2_merge_next_source_mismatch",
        [("前段  ", None), ("目前後段", None)],
        {"paragraph_merge_candidates": [_merge_candidate()]},
    )

    assert _texts(_read_docx(result.output_path)) == ["前段  ", "目前後段"]
    assert _result_codes(result) == ["PARAGRAPH_MERGE_SOURCE_MISMATCH"]


def test_conflicting_merge_candidate_is_skipped() -> None:
    result = _apply_payload(
        "phase2e2_merge_conflict",
        [("前段  ", None), ("  後段", None)],
        {
            "paragraph_merge_candidates": [
                _merge_candidate(candidate_id="first"),
                _merge_candidate(candidate_id="second"),
            ]
        },
    )

    assert _texts(_read_docx(result.output_path)) == ["前段後段"]
    assert _result_codes(result) == ["APPLIED_PARAGRAPH_MERGE", "PARAGRAPH_MERGE_CONFLICT"]
    assert result.apply_result.applied_count == 1
    assert result.apply_result.skipped_count == 1


def test_multiple_merge_candidates_are_processed_in_descending_index_order() -> None:
    result = _apply_payload(
        "phase2e2_merge_descending_order",
        [("A", None), ("B", None), ("C", None), ("D", None)],
        {
            "paragraph_merge_candidates": [
                _merge_candidate(
                    candidate_id="low",
                    paragraph_index=0,
                    next_paragraph_index=1,
                    source_text="A",
                    next_source_text="B",
                ),
                _merge_candidate(
                    candidate_id="high",
                    paragraph_index=2,
                    next_paragraph_index=3,
                    source_text="C",
                    next_source_text="D",
                ),
            ]
        },
    )

    assert _texts(_read_docx(result.output_path)) == ["AB", "CD"]
    assert _result_codes(result) == ["APPLIED_PARAGRAPH_MERGE", "APPLIED_PARAGRAPH_MERGE"]


def test_heading_boundary_is_not_merged() -> None:
    result = _apply_payload(
        "phase2e2_merge_heading_boundary",
        [("第一章", "Heading 2"), ("正文", None)],
        {
            "paragraph_merge_candidates": [
                _merge_candidate(source_text="第一章", next_source_text="正文"),
            ]
        },
    )

    document = _read_docx(result.output_path)
    assert _texts(document) == ["第一章", "正文"]
    assert document.paragraphs[0].style.name == "Heading 2"
    assert _result_codes(result) == ["PARAGRAPH_MERGE_CHAPTER_BOUNDARY"]


def test_toc_is_inserted_after_paragraph_merge_smoke() -> None:
    result = _apply_payload(
        "phase2e2_merge_before_toc",
        [("第一章", "Heading 2"), ("前段  ", None), ("  後段", None)],
        {
            "paragraph_merge_candidates": [
                _merge_candidate(paragraph_index=1, next_paragraph_index=2),
            ]
        },
        create_toc=True,
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_FIELD_INSERTED
    document = _read_docx(result.output_path)
    assert _texts(document)[:5] == ["目錄", "請在 Word 中更新目錄。", "第一章", "前段後段"]
    assert _result_codes(result) == ["APPLIED_PARAGRAPH_MERGE"]
