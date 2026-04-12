from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, convert
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _read_docx_text(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def _write_phase1c_config(tmp: Path) -> Path:
    low_path = tmp / "low.yaml"
    high_path = tmp / "high.yaml"
    config_path = tmp / "config.yaml"
    low_path.write_text(
        '- source: "火車"\n'
        '  target: "列車"\n'
        "  risk_level: low\n"
        "  enabled: true\n"
        "  category: test\n"
        '- source: "資訊"\n'
        '  target: "情報"\n'
        "  risk_level: low\n"
        "  enabled: true\n"
        "  category: test\n",
        encoding="utf-8",
    )
    high_path.write_text(
        '- term: "資訊"\n'
        "  risk_category: wording\n"
        "  enabled: true\n",
        encoding="utf-8",
    )
    config_path.write_text(
        "opencc_config: s2t\n"
        "default_report_name: report.xlsx\n"
        f"default_term_dict_path: {low_path.as_posix()}\n"
        f"default_high_risk_rules_path: {high_path.as_posix()}\n"
        "enable_space_cleanup: true\n"
        "document_format:\n"
        "  page_margin_top_cm: 1.27\n"
        "  page_margin_bottom_cm: 1.27\n"
        "  page_margin_left_cm: 1.27\n"
        "  page_margin_right_cm: 1.27\n"
        "  body_first_line_indent_chars: 2\n"
        "  heading_style_name: Heading 2\n",
        encoding="utf-8",
    )
    return config_path


def test_convert_outputs_docx_without_modifying_input() -> None:
    tmp = make_test_dir("phase1c_output")
    input_path = tmp / "novel.docx"
    output_dir = tmp / "out"
    config_path = _write_phase1c_config(tmp)
    _make_docx(input_path, ["软件和火车"])
    before_size = input_path.stat().st_size
    before_mtime = input_path.stat().st_mtime_ns

    result = convert(Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path))

    assert result.operation == "convert"
    assert result.docx_processed is True
    assert result.output_path == output_dir / "novel_converted.docx"
    assert result.output_path.exists()
    assert result.schema.errors == []
    assert result.schema.warnings == []
    assert result.schema.paragraph_merge_candidates == []
    assert result.config_check.warnings
    assert input_path.stat().st_size == before_size
    assert input_path.stat().st_mtime_ns == before_mtime


def test_convert_applies_opencc_and_low_risk_terms() -> None:
    tmp = make_test_dir("phase1c_low_risk")
    input_path = tmp / "terms.docx"
    output_dir = tmp / "out"
    config_path = _write_phase1c_config(tmp)
    _make_docx(input_path, ["火车"])

    result = convert(Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path))

    assert result.output_path is not None
    assert _read_docx_text(result.output_path) == ["列車"]


def test_convert_keeps_high_risk_paragraph_to_opencc_only() -> None:
    tmp = make_test_dir("phase1c_high_risk")
    input_path = tmp / "risk.docx"
    output_dir = tmp / "out"
    config_path = _write_phase1c_config(tmp)
    _make_docx(input_path, ["資訊", "火车"])

    result = convert(Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path))

    assert result.output_path is not None
    assert _read_docx_text(result.output_path) == ["資訊", "列車"]
    assert len(result.schema.chapter_candidates) == 0
    assert len(result.schema.review_candidates) == 1
    assert result.schema.review_candidates[0].source_text == "資訊"


def test_convert_keeps_analysis_chapter_candidates() -> None:
    tmp = make_test_dir("phase1c_keep_chapters")
    input_path = tmp / "chapter.docx"
    output_dir = tmp / "out"
    config_path = _write_phase1c_config(tmp)
    _make_docx(input_path, ["第十章 夜雨", "火车"])

    result = convert(Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path))

    assert result.schema.errors == []
    assert len(result.schema.chapter_candidates) == 1
    chapter = result.schema.chapter_candidates[0]
    assert chapter.status == "pending"
    assert chapter.auto_accept is False


def test_convert_requires_output_dir() -> None:
    tmp = make_test_dir("phase1c_no_output_dir")
    input_path = tmp / "input.docx"
    config_path = _write_phase1c_config(tmp)
    _make_docx(input_path, ["火车"])

    result = convert(Phase1Options(input_path=input_path, config_path=config_path))

    assert result.docx_processed is False
    assert result.output_path is None
    assert result.schema.errors
    assert result.schema.errors[0].code == "OUTPUT_DIR_REQUIRED"


def test_convert_missing_input_path_uses_existing_input_not_found_code() -> None:
    tmp = make_test_dir("phase1c_no_input")
    config_path = _write_phase1c_config(tmp)

    result = convert(Phase1Options(output_dir=tmp / "out", config_path=config_path))

    assert result.docx_processed is False
    assert result.output_path is None
    assert result.schema.errors
    assert result.schema.errors[0].code == "INPUT_NOT_FOUND"


def test_convert_uses_conflict_safe_output_name() -> None:
    tmp = make_test_dir("phase1c_conflict")
    input_path = tmp / "novel.docx"
    output_dir = tmp / "out"
    output_dir.mkdir()
    config_path = _write_phase1c_config(tmp)
    _make_docx(input_path, ["火车"])
    _make_docx(output_dir / "novel_converted.docx", ["existing"])

    result = convert(Phase1Options(input_path=input_path, output_dir=output_dir, config_path=config_path))

    assert result.output_path == output_dir / "novel_converted_001.docx"
    assert result.output_path.exists()
