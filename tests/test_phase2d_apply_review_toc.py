import json
from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, apply_review
from src.phase2_toc_builder import (
    TOC_STATUS_FALLBACK_CHAPTER_LIST,
    TOC_STATUS_FIELD_INSERTED,
    TOC_STATUS_NOT_REQUESTED,
    TOC_STATUS_SKIPPED_EXISTING_TOC,
    TOC_STATUS_SKIPPED_NO_HEADINGS,
)
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[tuple[str, str | None]]) -> None:
    document = Document()
    for text, style_name in paragraphs:
        paragraph = document.add_paragraph(text)
        if style_name is not None:
            paragraph.style = style_name
    document.save(path)


def _read_docx(path: Path) -> Document:
    return Document(path)


def _write_payload(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _chapter_candidate(**overrides) -> dict:
    data = {
        "candidate_id": "chapter:0",
        "type": "chapter",
        "status": "accepted",
        "paragraph_index": 0,
    }
    data.update(overrides)
    return data


def _review_candidate(**overrides) -> dict:
    data = {
        "candidate_id": "risk:one",
        "type": "high_risk_term",
        "status": "accepted",
        "source_text": "bad",
        "resolved_text": "good",
        "paragraph_index": 1,
        "char_start": 0,
        "char_end": 3,
        "context_before": "",
        "context_after": "",
    }
    data.update(overrides)
    return data


def _apply_payload(
    name: str,
    paragraphs: list[tuple[str, str | None]],
    payload: dict,
    create_toc: bool = True,
):
    tmp = make_test_dir(name)
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, paragraphs)
    _write_payload(review_path, payload)
    return apply_review(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            apply_review_path=review_path,
            create_toc=create_toc,
        )
    )


def _texts(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def test_apply_review_inserts_toc_after_chapter_candidate_apply() -> None:
    result = _apply_payload(
        "phase2d_apply_review_toc",
        [("Chapter 1", None), ("Body", None)],
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_FIELD_INSERTED
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert _texts(document)[:4] == ["目錄", "請在 Word 中更新目錄。", "Chapter 1", "Body"]
    assert document.paragraphs[2].style.name == "Heading 2"


def test_apply_review_respects_no_create_toc() -> None:
    result = _apply_payload(
        "phase2d_apply_review_no_toc",
        [("Chapter 1", None), ("Body", None)],
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
        create_toc=False,
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_NOT_REQUESTED
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert _texts(document) == ["Chapter 1", "Body"]
    assert document.paragraphs[0].style.name == "Heading 2"


def test_apply_review_skips_toc_without_heading_2() -> None:
    result = _apply_payload(
        "phase2d_apply_review_no_heading",
        [("Chapter 1", None), ("Body", None)],
        {},
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_SKIPPED_NO_HEADINGS
    assert result.output_path is not None
    assert _texts(_read_docx(result.output_path)) == ["Chapter 1", "Body"]


def test_apply_review_skips_existing_toc_without_duplicate_toc_block() -> None:
    result = _apply_payload(
        "phase2d_apply_review_existing_toc",
        [("目錄", None), ("Chapter 1", "Heading 2"), ("Body", None)],
        {},
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_SKIPPED_EXISTING_TOC
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert _texts(document).count("目錄") == 1
    assert _texts(document) == ["目錄", "Chapter 1", "Body"]


def test_apply_review_toc_field_failure_falls_back_without_schema_error(monkeypatch) -> None:
    def fail_field_insert(document):
        raise RuntimeError("field failed")

    monkeypatch.setattr("src.phase2_toc_builder._insert_toc_field", fail_field_insert)
    result = _apply_payload(
        "phase2d_apply_review_toc_fallback",
        [("Chapter 1", None), ("Body", None)],
        {"chapter_candidates": [_chapter_candidate(paragraph_index=0)]},
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_FALLBACK_CHAPTER_LIST
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert _texts(document)[:3] == ["目錄", "Chapter 1", "Chapter 1"]


def test_apply_review_keeps_review_candidate_and_chapter_candidate_behavior_with_toc() -> None:
    result = _apply_payload(
        "phase2d_apply_review_combined",
        [("Chapter 1", None), ("bad body", None)],
        {
            "chapter_candidates": [_chapter_candidate(paragraph_index=0)],
            "review_candidates": [_review_candidate()],
        },
    )

    assert result.schema.errors == []
    assert result.schema.toc.status == TOC_STATUS_FIELD_INSERTED
    assert result.output_path is not None
    document = _read_docx(result.output_path)
    assert _texts(document)[:4] == ["目錄", "請在 Word 中更新目錄。", "Chapter 1", "good body"]
    assert document.paragraphs[2].style.name == "Heading 2"
    assert {item.result_code for item in result.apply_result.candidate_results} == {
        "APPLIED",
        "APPLIED_CHAPTER_HEADING",
    }
