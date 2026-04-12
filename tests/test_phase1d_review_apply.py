import json
from pathlib import Path

from docx import Document

from src.phase1_pipeline import Phase1Options, apply_review
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _read_docx_text(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def _write_review(path: Path, candidates: list[dict]) -> None:
    path.write_text(json.dumps({"review_candidates": candidates}, ensure_ascii=False), encoding="utf-8")


def _candidate(**overrides):
    data = {
        "candidate_id": "risk:one",
        "type": "high_risk_term",
        "status": "accepted",
        "source_text": "資訊",
        "resolved_text": "訊息",
        "paragraph_index": 0,
        "char_start": 0,
        "char_end": 2,
        "context_before": "",
        "context_after": "",
    }
    data.update(overrides)
    return data


def test_apply_review_applies_accepted_candidate_without_context() -> None:
    tmp = make_test_dir("phase1d_apply")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊"])
    _write_review(review_path, [_candidate()])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert result.docx_processed is True
    assert result.output_path == output_dir / "source_reviewed.docx"
    assert _read_docx_text(result.output_path) == ["訊息"]
    assert result.apply_result is not None
    assert result.apply_result.applied_count == 1
    assert result.apply_result.candidate_results[0].result_code == "APPLIED"
    assert result.schema.warnings == []


def test_apply_review_skips_disallowed_statuses() -> None:
    tmp = make_test_dir("phase1d_status")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊資訊資訊資訊"])
    _write_review(
        review_path,
        [
            _candidate(candidate_id="pending", status="pending", char_start=0, char_end=2),
            _candidate(candidate_id="rejected", status="rejected", char_start=2, char_end=4),
            _candidate(candidate_id="skip", status="skip", char_start=4, char_end=6),
        ],
    )

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert _read_docx_text(result.output_path) == ["資訊資訊資訊資訊"]
    assert {item.result_code for item in result.apply_result.candidate_results} == {"SKIPPED_STATUS"}


def test_apply_review_requires_non_empty_resolved_text() -> None:
    tmp = make_test_dir("phase1d_empty_resolved")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊"])
    _write_review(review_path, [_candidate(resolved_text="   ")])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert _read_docx_text(result.output_path) == ["資訊"]
    assert result.apply_result.candidate_results[0].result_code == "SKIPPED_EMPTY_RESOLVED_TEXT"


def test_apply_review_records_unsupported_type() -> None:
    tmp = make_test_dir("phase1d_unsupported_type")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊"])
    _write_review(review_path, [_candidate(type="chapter")])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert _read_docx_text(result.output_path) == ["資訊"]
    assert result.apply_result.candidate_results[0].result_code == "SKIPPED_UNSUPPORTED_TYPE"


def test_apply_review_context_mismatch_is_not_applied() -> None:
    tmp = make_test_dir("phase1d_context_mismatch")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["前文資訊後文"])
    _write_review(review_path, [_candidate(char_start=2, char_end=4, context_before="不符")])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert _read_docx_text(result.output_path) == ["前文資訊後文"]
    assert result.apply_result.candidate_results[0].result_code == "APPLY_TARGET_NOT_FOUND"


def test_apply_review_span_invalid_is_not_applied() -> None:
    tmp = make_test_dir("phase1d_span_invalid")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊"])
    _write_review(review_path, [_candidate(char_start=0, char_end=20)])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert _read_docx_text(result.output_path) == ["資訊"]
    assert result.apply_result.candidate_results[0].result_code == "REVIEW_SPAN_INVALID"


def test_apply_review_applies_same_paragraph_candidates_from_right_to_left() -> None:
    tmp = make_test_dir("phase1d_right_to_left")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊資訊"])
    _write_review(
        review_path,
        [
            _candidate(candidate_id="left", char_start=0, char_end=2, resolved_text="訊息"),
            _candidate(candidate_id="right", char_start=2, char_end=4, resolved_text="消息"),
        ],
    )

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert _read_docx_text(result.output_path) == ["訊息消息"]
    assert [item.result_code for item in result.apply_result.candidate_results] == ["APPLIED", "APPLIED"]


def test_apply_review_conflict_is_skipped() -> None:
    tmp = make_test_dir("phase1d_conflict")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    _make_docx(input_path, ["資訊測試"])
    _write_review(
        review_path,
        [
            _candidate(candidate_id="wide", source_text="資訊測試", char_start=0, char_end=4, resolved_text="內容"),
            _candidate(candidate_id="narrow", char_start=0, char_end=2, resolved_text="訊息"),
        ],
    )

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert any(item.result_code == "APPLY_CONFLICT" for item in result.apply_result.candidate_results)


def test_apply_review_output_name_avoids_conflict() -> None:
    tmp = make_test_dir("phase1d_name_conflict")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    output_dir.mkdir()
    _make_docx(input_path, ["資訊"])
    _make_docx(output_dir / "source_reviewed.docx", ["existing"])
    _write_review(review_path, [_candidate()])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=output_dir, apply_review_path=review_path))

    assert result.output_path == output_dir / "source_reviewed_001.docx"


def test_apply_review_uses_exact_reviewed_output_path() -> None:
    tmp = make_test_dir("phase1d_exact_output")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    reviewed_output = tmp / "custom" / "exact.docx"
    _make_docx(input_path, [_candidate()["source_text"]])
    _write_review(review_path, [_candidate()])

    result = apply_review(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            apply_review_path=review_path,
            reviewed_output_path=reviewed_output,
        )
    )

    assert result.output_path == reviewed_output.resolve()
    assert reviewed_output.exists()
    assert not output_dir.exists()
    assert result.apply_result.candidate_results[0].result_code == "APPLIED"
    assert _read_docx_text(reviewed_output) == [_candidate()["resolved_text"]]


def test_apply_review_relative_exact_output_path_uses_current_working_directory(monkeypatch) -> None:
    tmp = make_test_dir("phase1d_exact_relative")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    _make_docx(input_path, [_candidate()["source_text"]])
    _write_review(review_path, [_candidate()])
    monkeypatch.chdir(tmp)

    result = apply_review(
        Phase1Options(
            input_path=input_path,
            apply_review_path=review_path,
            reviewed_output_path=Path("relative_reviewed.docx"),
        )
    )

    assert result.output_path == (tmp / "relative_reviewed.docx").resolve()
    assert (tmp / "relative_reviewed.docx").exists()


def test_apply_review_exact_output_path_takes_priority_over_output_dir() -> None:
    tmp = make_test_dir("phase1d_exact_priority")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    output_dir = tmp / "out"
    reviewed_output = tmp / "elsewhere" / "source_reviewed.docx"
    output_dir.mkdir()
    _make_docx(output_dir / "source_reviewed.docx", ["existing"])
    _make_docx(input_path, [_candidate()["source_text"]])
    _write_review(review_path, [_candidate()])

    result = apply_review(
        Phase1Options(
            input_path=input_path,
            output_dir=output_dir,
            apply_review_path=review_path,
            reviewed_output_path=reviewed_output,
        )
    )

    assert result.output_path == reviewed_output.resolve()
    assert _read_docx_text(output_dir / "source_reviewed.docx") == ["existing"]


def test_apply_review_exact_output_existing_file_is_error() -> None:
    tmp = make_test_dir("phase1d_exact_exists")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    reviewed_output = tmp / "already.docx"
    _make_docx(input_path, [_candidate()["source_text"]])
    _make_docx(reviewed_output, ["existing"])
    _write_review(review_path, [_candidate()])

    result = apply_review(
        Phase1Options(input_path=input_path, apply_review_path=review_path, reviewed_output_path=reviewed_output)
    )

    assert result.schema.errors[0].code == "REVIEWED_OUTPUT_EXISTS"
    assert _read_docx_text(reviewed_output) == ["existing"]


def test_apply_review_exact_output_requires_docx_suffix() -> None:
    tmp = make_test_dir("phase1d_exact_not_docx")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    _make_docx(input_path, [_candidate()["source_text"]])
    _write_review(review_path, [_candidate()])

    result = apply_review(
        Phase1Options(input_path=input_path, apply_review_path=review_path, reviewed_output_path=tmp / "reviewed.txt")
    )

    assert result.schema.errors[0].code == "REVIEWED_OUTPUT_NOT_DOCX"


def test_apply_review_exact_output_directory_is_error() -> None:
    tmp = make_test_dir("phase1d_exact_directory")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    reviewed_output = tmp / "directory.docx"
    reviewed_output.mkdir()
    _make_docx(input_path, [_candidate()["source_text"]])
    _write_review(review_path, [_candidate()])

    result = apply_review(
        Phase1Options(input_path=input_path, apply_review_path=review_path, reviewed_output_path=reviewed_output)
    )

    assert result.schema.errors[0].code == "REVIEWED_OUTPUT_IS_DIRECTORY"


def test_apply_review_exact_output_parent_create_failure(monkeypatch) -> None:
    tmp = make_test_dir("phase1d_exact_parent_create_failed")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    reviewed_output = tmp / "blocked" / "reviewed.docx"
    _make_docx(input_path, [_candidate()["source_text"]])
    _write_review(review_path, [_candidate()])
    original_mkdir = Path.mkdir

    def fake_mkdir(self, *args, **kwargs):
        if self == reviewed_output.parent:
            raise OSError("blocked")
        return original_mkdir(self, *args, **kwargs)

    monkeypatch.setattr(Path, "mkdir", fake_mkdir)

    result = apply_review(
        Phase1Options(input_path=input_path, apply_review_path=review_path, reviewed_output_path=reviewed_output)
    )

    assert result.schema.errors[0].code == "REVIEWED_OUTPUT_DIR_CREATE_FAILED"


def test_apply_review_invalid_json_returns_schema_error() -> None:
    tmp = make_test_dir("phase1d_invalid_json")
    input_path = tmp / "source.docx"
    review_path = tmp / "review.json"
    _make_docx(input_path, ["資訊"])
    review_path.write_text("{", encoding="utf-8")

    result = apply_review(Phase1Options(input_path=input_path, output_dir=tmp / "out", apply_review_path=review_path))

    assert result.schema.errors[0].code == "REVIEW_JSON_INVALID"


def test_apply_review_missing_json_returns_schema_error() -> None:
    tmp = make_test_dir("phase1d_missing_json")
    input_path = tmp / "source.docx"
    _make_docx(input_path, ["資訊"])

    result = apply_review(Phase1Options(input_path=input_path, output_dir=tmp / "out", apply_review_path=tmp / "missing.json"))

    assert result.schema.errors[0].code == "REVIEW_JSON_NOT_FOUND"
