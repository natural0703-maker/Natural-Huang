from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from src.processing_service_v3 import ApplyReviewOptions, ProcessingOptions, run_apply_review, run_processing
from src.report_writer import save_workbook_safely
from tests.test_paths import make_test_dir


def _make_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _set_review_rows(review_summary_path: Path, updates: dict[str, tuple[str, str, int | None]]) -> None:
    wb = load_workbook(review_summary_path)
    ws = wb["review_summary"]
    headers = [cell.value for cell in ws[1]]
    idx = {str(v): i for i, v in enumerate(headers)}
    for row in ws.iter_rows(min_row=2):
        hit_term = str(row[idx["hit_term"]].value or "")
        if hit_term not in updates:
            continue
        status, resolved_text, paragraph_index = updates[hit_term]
        row[idx["status"]].value = status
        row[idx["resolved_text"]].value = resolved_text
        if paragraph_index is not None:
            row[idx["paragraph_index"]].value = paragraph_index
    save_workbook_safely(wb, review_summary_path)


def test_apply_review_only_accept_and_non_empty_resolved_text() -> None:
    tmp = make_test_dir("apply_review_accept_only")
    src_doc = tmp / "novel.docx"
    out_dir = tmp / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    _make_docx(src_doc, ["這裡的資訊支持測試。"])

    process_result = run_processing(ProcessingOptions(input_file=src_doc, output_dir=out_dir))
    first_output = Path(process_result.batch_result.summaries[0].output_file)
    review_summary = process_result.review_summary_path

    _set_review_rows(
        review_summary,
        {
            "支持": ("accepted", "支援", None),
            "資訊": ("rejected", "訊息", None),
        },
    )

    apply_result = run_apply_review(
        ApplyReviewOptions(
            input_file=first_output,
            output_dir=out_dir,
            apply_review_summary_path=review_summary,
        )
    )
    assert apply_result.applied_count == 1

    reviewed_output = apply_result.output_files[0]
    reviewed_doc = Document(str(reviewed_output))
    reviewed_text = reviewed_doc.paragraphs[0].text
    assert "支援" in reviewed_text
    assert "資訊" in reviewed_text

    original_doc = Document(str(first_output))
    assert "支持" in original_doc.paragraphs[0].text


def test_apply_review_skips_pending_rejected_skip() -> None:
    tmp = make_test_dir("apply_review_skip_status")
    src_doc = tmp / "novel.docx"
    out_dir = tmp / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    _make_docx(src_doc, ["資訊支持影片。"])

    process_result = run_processing(ProcessingOptions(input_file=src_doc, output_dir=out_dir))
    first_output = Path(process_result.batch_result.summaries[0].output_file)
    review_summary = process_result.review_summary_path

    _set_review_rows(
        review_summary,
        {
            "資訊": ("pending", "訊息", None),
            "支持": ("rejected", "支援", None),
            "影片": ("skip", "視訊", None),
        },
    )

    apply_result = run_apply_review(
        ApplyReviewOptions(
            input_file=first_output,
            output_dir=out_dir,
            apply_review_summary_path=review_summary,
        )
    )
    assert apply_result.applied_count == 0
    assert apply_result.skipped_count >= 3


def test_apply_review_not_found_is_recorded_and_summary_exists() -> None:
    tmp = make_test_dir("apply_review_not_found")
    src_doc = tmp / "novel.docx"
    out_dir = tmp / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    _make_docx(src_doc, ["資訊支持測試。"])

    process_result = run_processing(ProcessingOptions(input_file=src_doc, output_dir=out_dir))
    first_output = Path(process_result.batch_result.summaries[0].output_file)
    review_summary = process_result.review_summary_path

    _set_review_rows(
        review_summary,
        {
            "資訊": ("accepted", "訊息", 999),
        },
    )

    apply_result = run_apply_review(
        ApplyReviewOptions(
            input_file=first_output,
            output_dir=out_dir,
            apply_review_summary_path=review_summary,
        )
    )

    assert apply_result.not_found_count >= 1
    assert apply_result.apply_summary_path.exists()
    wb = load_workbook(apply_result.apply_summary_path)
    assert "apply_summary" in wb.sheetnames
    assert "apply_details" in wb.sheetnames
    assert "apply_failures" in wb.sheetnames
    assert "apply_reason_stats" in wb.sheetnames


def test_apply_review_multiple_accepted_in_same_paragraph() -> None:
    tmp = make_test_dir("apply_review_multi_accepted")
    src_doc = tmp / "novel.docx"
    out_dir = tmp / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    _make_docx(src_doc, ["資訊支持影片。"])

    process_result = run_processing(ProcessingOptions(input_file=src_doc, output_dir=out_dir))
    first_output = Path(process_result.batch_result.summaries[0].output_file)
    review_summary = process_result.review_summary_path

    _set_review_rows(
        review_summary,
        {
            "資訊": ("accepted", "訊息", None),
            "支持": ("accepted", "支援", None),
        },
    )

    apply_result = run_apply_review(
        ApplyReviewOptions(
            input_file=first_output,
            output_dir=out_dir,
            apply_review_summary_path=review_summary,
        )
    )
    assert apply_result.applied_count >= 2
    reviewed_doc = Document(str(apply_result.output_files[0]))
    reviewed_text = reviewed_doc.paragraphs[0].text
    assert "訊息" in reviewed_text
    assert "支援" in reviewed_text
