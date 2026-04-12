from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from src.config_loader import DocumentFormatConfig


@dataclass(frozen=True)
class ParagraphOutput:
    text: str
    is_heading: bool = False


DEFAULT_FORMAT = DocumentFormatConfig()


def build_output_docx_path(input_path: Path, output_dir: Path) -> Path:
    stem = input_path.stem
    base_name = f"{stem}_TW.docx"
    candidate = output_dir / base_name
    if not candidate.exists():
        return candidate

    index = 1
    while True:
        candidate = output_dir / f"{stem}_TW_{index:03d}.docx"
        if not candidate.exists():
            return candidate
        index += 1


def build_reviewed_output_docx_path(input_path: Path, output_dir: Path) -> Path:
    stem = input_path.stem
    base_name = f"{stem}_reviewed.docx"
    candidate = output_dir / base_name
    if not candidate.exists():
        return candidate

    index = 1
    while True:
        candidate = output_dir / f"{stem}_reviewed_{index:03d}.docx"
        if not candidate.exists():
            return candidate
        index += 1


def _resolve_style_name(document: Document, candidates: list[str]) -> str | None:
    existing = {style.name for style in document.styles}
    for candidate in candidates:
        if candidate in existing:
            return candidate
    return None


def _apply_run_font(run, font_name: str, font_size_pt: float) -> None:  # type: ignore[no-untyped-def]
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def _normalize_paragraphs(paragraphs: list[str] | list[ParagraphOutput]) -> list[ParagraphOutput]:
    normalized: list[ParagraphOutput] = []
    for item in paragraphs:
        if isinstance(item, ParagraphOutput):
            normalized.append(item)
        else:
            normalized.append(ParagraphOutput(text=str(item), is_heading=False))
    return normalized


def _resolve_line_spacing(mode: str) -> WD_LINE_SPACING:
    normalized = mode.strip().lower()
    if normalized == "exact":
        return WD_LINE_SPACING.EXACTLY
    if normalized == "single":
        return WD_LINE_SPACING.SINGLE
    if normalized in {"one_point_five", "1.5"}:
        return WD_LINE_SPACING.ONE_POINT_FIVE
    if normalized == "double":
        return WD_LINE_SPACING.DOUBLE
    return WD_LINE_SPACING.AT_LEAST


def write_paragraphs_to_docx(
    paragraphs: list[str] | list[ParagraphOutput],
    output_path: Path,
    document_format: DocumentFormatConfig | None = None,
) -> None:
    format_cfg = document_format if document_format is not None else DEFAULT_FORMAT

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(format_cfg.page_margin_top_cm)
        section.bottom_margin = Cm(format_cfg.page_margin_bottom_cm)
        section.left_margin = Cm(format_cfg.page_margin_left_cm)
        section.right_margin = Cm(format_cfg.page_margin_right_cm)

    heading_style = _resolve_style_name(
        document,
        [format_cfg.heading_style_name, "Heading 2", "\u6a19\u984c 2"],
    )
    body_style = _resolve_style_name(document, ["\u5167\u6587", "\u672c\u6587", "Normal", "\u6a19\u6e96"])

    for item in _normalize_paragraphs(paragraphs):
        paragraph = document.add_paragraph(item.text)

        if item.is_heading and heading_style is not None:
            paragraph.style = heading_style
            continue

        if body_style is not None:
            paragraph.style = body_style

        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(format_cfg.body_space_before_pt)
        paragraph_format.space_after = Pt(format_cfg.body_space_after_pt)
        paragraph_format.line_spacing_rule = _resolve_line_spacing(format_cfg.body_line_spacing_mode)
        paragraph_format.line_spacing = Pt(format_cfg.body_min_line_height_pt)
        paragraph_format.first_line_indent = Pt(
            format_cfg.body_first_line_indent_chars * format_cfg.body_font_size_pt
        )

        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            _apply_run_font(run, format_cfg.body_font_name, format_cfg.body_font_size_pt)

    document.save(str(output_path))
