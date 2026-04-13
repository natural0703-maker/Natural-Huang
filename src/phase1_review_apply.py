from __future__ import annotations

import json
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from src.frozen_spec_v1 import HEADING_STYLE_NAME
from src.phase2_toc_builder import TOC_STATUS_FAILED, insert_minimal_toc
from src.review_schema import ErrorRecord, ReviewSchema, validate_reviewed_json_payload


PARAGRAPH_MERGE_APPLIED_CODES = frozenset({"APPLIED_PARAGRAPH_MERGE"})
PARAGRAPH_MERGE_SKIPPED_CODES = frozenset(
    {
        "SKIPPED_PARAGRAPH_MERGE_STATUS",
        "SKIPPED_UNSUPPORTED_TYPE",
        "PARAGRAPH_MERGE_CONFLICT",
        "PARAGRAPH_MERGE_CHAPTER_BOUNDARY",
    }
)
PARAGRAPH_MERGE_FAILED_CODES = frozenset(
    {
        "PARAGRAPH_MERGE_PARAGRAPH_INDEX_INVALID",
        "PARAGRAPH_MERGE_NOT_ADJACENT",
        "PARAGRAPH_MERGE_PARAGRAPH_EMPTY",
        "PARAGRAPH_MERGE_SOURCE_MISMATCH",
        "PARAGRAPH_MERGE_APPLY_FAILED",
    }
)
PARAGRAPH_MERGE_RESULT_CODES = (
    PARAGRAPH_MERGE_APPLIED_CODES | PARAGRAPH_MERGE_SKIPPED_CODES | PARAGRAPH_MERGE_FAILED_CODES
)
PARAGRAPH_MERGE_SOURCE_TEXT_MISMATCH = "SOURCE_TEXT_MISMATCH"
PARAGRAPH_MERGE_NEXT_SOURCE_TEXT_MISMATCH = "NEXT_SOURCE_TEXT_MISMATCH"
PARAGRAPH_MERGE_DIAGNOSTIC_SAMPLE_LIMIT = 3


@dataclass(frozen=True)
class ReviewApplyCandidateResult:
    candidate_id: str
    status: str
    result_code: str
    paragraph_index: int | None
    applied: bool
    message: str
    detail_code: str = ""


@dataclass(frozen=True)
class ParagraphMergeSummary:
    applied_count: int = 0
    skipped_count: int = 0
    failed_count: int = 0
    codes: dict[str, int] = field(default_factory=dict)


@dataclass(frozen=True)
class ParagraphMergeDiagnostics:
    source_mismatch_count: int = 0
    next_source_mismatch_count: int = 0
    total_mismatch_count: int = 0
    sample_candidate_ids: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class ReviewApplyResult:
    schema: ReviewSchema
    output_path: Path | None
    applied: bool
    candidate_results: list[ReviewApplyCandidateResult]
    applied_count: int
    skipped_count: int
    failed_count: int
    paragraph_merge_summary: ParagraphMergeSummary = field(default_factory=ParagraphMergeSummary)
    paragraph_merge_diagnostics: ParagraphMergeDiagnostics = field(default_factory=ParagraphMergeDiagnostics)


def apply_review_docx(
    input_path: Path | None,
    review_json_path: Path | None,
    output_dir: Path | None,
    reviewed_output_path: Path | None = None,
    create_toc: bool = True,
) -> ReviewApplyResult:
    if input_path is None or not input_path.exists():
        return _flow_error("INPUT_NOT_FOUND", "找不到輸入 DOCX。", str(input_path or ""))
    if input_path.suffix.lower() != ".docx":
        return _flow_error("INPUT_NOT_DOCX", f"輸入檔不是 .docx：{input_path}", str(input_path))
    if review_json_path is None or not review_json_path.exists():
        return _flow_error("REVIEW_JSON_NOT_FOUND", "找不到 reviewed JSON。", str(review_json_path or ""))
    output_path_result = _reviewed_output_path(reviewed_output_path, output_dir, input_path.stem)
    if isinstance(output_path_result, ReviewApplyResult):
        return output_path_result
    output_path = output_path_result

    try:
        raw = json.loads(review_json_path.read_text(encoding="utf-8"))
    except Exception as exc:
        return _flow_error("REVIEW_JSON_INVALID", f"無法讀取 reviewed JSON：{review_json_path}", str(exc))
    review_json_errors = validate_reviewed_json_payload(raw)
    if review_json_errors:
        return _flow_error_record(review_json_errors[0])
    raw_chapter_candidates = raw.get("chapter_candidates", [])
    raw_candidates = raw.get("review_candidates", [])
    raw_paragraph_merge_candidates = raw.get("paragraph_merge_candidates", [])

    try:
        document = Document(input_path)
    except Exception as exc:
        return _flow_error("DOCX_READ_FAILED", f"無法讀取 DOCX：{input_path}", str(exc))

    candidate_results: list[ReviewApplyCandidateResult] = []
    applicable: list[dict[str, Any]] = []
    for item in raw_candidates:
        candidate, result = _validate_candidate(item)
        if result is not None:
            candidate_results.append(result)
            continue
        assert candidate is not None
        applicable.append(candidate)

    paragraphs = document.paragraphs
    occupied_ranges: dict[int, list[tuple[int, int]]] = {}
    for candidate in sorted(
        applicable,
        key=lambda item: (item["paragraph_index"], item["char_start"], item["char_end"]),
        reverse=True,
    ):
        paragraph_index = candidate["paragraph_index"]
        char_start = candidate["char_start"]
        char_end = candidate["char_end"]
        conflict = _has_conflict(occupied_ranges.get(paragraph_index, []), char_start, char_end)
        if conflict:
            candidate_results.append(_candidate_result(candidate, "APPLY_CONFLICT", False, "候選區間與其他候選重疊。"))
            continue
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            candidate_results.append(_candidate_result(candidate, "APPLY_TARGET_NOT_FOUND", False, "找不到目標段落。"))
            continue

        text = paragraphs[paragraph_index].text
        if char_start < 0 or char_end <= char_start or char_end > len(text):
            candidate_results.append(_candidate_result(candidate, "REVIEW_SPAN_INVALID", False, "候選 span 超出段落範圍。"))
            continue
        if text[char_start:char_end] != candidate["source_text"]:
            candidate_results.append(_candidate_result(candidate, "APPLY_TARGET_NOT_FOUND", False, "目標文字與 source_text 不符。"))
            continue
        context_before = candidate["context_before"]
        context_after = candidate["context_after"]
        if context_before and not text[:char_start].endswith(context_before):
            candidate_results.append(_candidate_result(candidate, "APPLY_TARGET_NOT_FOUND", False, "context_before 不符。"))
            continue
        if context_after and not text[char_end:].startswith(context_after):
            candidate_results.append(_candidate_result(candidate, "APPLY_TARGET_NOT_FOUND", False, "context_after 不符。"))
            continue

        paragraphs[paragraph_index].text = text[:char_start] + candidate["resolved_text"] + text[char_end:]
        occupied_ranges.setdefault(paragraph_index, []).append((char_start, char_end))
        candidate_results.append(_candidate_result(candidate, "APPLIED", True, "已套用。"))

    _apply_chapter_candidates(document, raw_chapter_candidates, candidate_results)
    _apply_paragraph_merge_candidates(document, raw_paragraph_merge_candidates, candidate_results)
    toc = insert_minimal_toc(document, requested=create_toc)
    if toc.status == TOC_STATUS_FAILED:
        return _flow_error("TOC_INSERT_FAILED", "無法建立目錄，也無法建立 fallback 章節清單。", "")

    try:
        document.save(output_path)
    except Exception as exc:
        if create_toc:
            return _flow_error("TOC_INSERT_FAILED", f"無法寫入含目錄的 reviewed DOCX：{output_path}", str(exc))
        return _flow_error("DOCX_WRITE_FAILED", f"無法寫入 reviewed DOCX：{output_path}", str(exc))

    applied_count = sum(1 for item in candidate_results if item.applied)
    skipped_count = sum(1 for item in candidate_results if not item.applied and _is_skipped_result(item.result_code))
    failed_count = sum(1 for item in candidate_results if not item.applied and not _is_skipped_result(item.result_code))
    paragraph_merge_summary = build_paragraph_merge_summary(candidate_results)
    paragraph_merge_diagnostics = build_paragraph_merge_diagnostics(candidate_results)
    return ReviewApplyResult(
        schema=ReviewSchema(toc=toc),
        output_path=output_path,
        applied=applied_count > 0,
        candidate_results=candidate_results,
        applied_count=applied_count,
        skipped_count=skipped_count,
        failed_count=failed_count,
        paragraph_merge_summary=paragraph_merge_summary,
        paragraph_merge_diagnostics=paragraph_merge_diagnostics,
    )


def build_paragraph_merge_summary(candidate_results: list[ReviewApplyCandidateResult]) -> ParagraphMergeSummary:
    codes = Counter(
        item.result_code for item in candidate_results if item.result_code in PARAGRAPH_MERGE_RESULT_CODES
    )
    return ParagraphMergeSummary(
        applied_count=sum(codes[code] for code in PARAGRAPH_MERGE_APPLIED_CODES),
        skipped_count=sum(codes[code] for code in PARAGRAPH_MERGE_SKIPPED_CODES),
        failed_count=sum(codes[code] for code in PARAGRAPH_MERGE_FAILED_CODES),
        codes=dict(sorted(codes.items())),
    )


def build_paragraph_merge_diagnostics(
    candidate_results: list[ReviewApplyCandidateResult],
) -> ParagraphMergeDiagnostics:
    source_mismatch_count = 0
    next_source_mismatch_count = 0
    sample_candidate_ids: list[str] = []
    for item in candidate_results:
        if item.result_code != "PARAGRAPH_MERGE_SOURCE_MISMATCH":
            continue
        detail_codes = set(filter(None, item.detail_code.split("|")))
        if PARAGRAPH_MERGE_SOURCE_TEXT_MISMATCH in detail_codes:
            source_mismatch_count += 1
        if PARAGRAPH_MERGE_NEXT_SOURCE_TEXT_MISMATCH in detail_codes:
            next_source_mismatch_count += 1
        if item.candidate_id and item.candidate_id not in sample_candidate_ids:
            if len(sample_candidate_ids) < PARAGRAPH_MERGE_DIAGNOSTIC_SAMPLE_LIMIT:
                sample_candidate_ids.append(item.candidate_id)
    return ParagraphMergeDiagnostics(
        source_mismatch_count=source_mismatch_count,
        next_source_mismatch_count=next_source_mismatch_count,
        total_mismatch_count=source_mismatch_count + next_source_mismatch_count,
        sample_candidate_ids=sample_candidate_ids,
    )


def _apply_chapter_candidates(
    document,
    raw_chapter_candidates: list[Any],
    candidate_results: list[ReviewApplyCandidateResult],
) -> None:
    paragraphs = document.paragraphs
    applied_paragraph_indices: set[int] = set()
    for item in raw_chapter_candidates:
        candidate, result = _validate_chapter_candidate(item)
        if result is not None:
            candidate_results.append(result)
            continue
        assert candidate is not None

        paragraph_index = candidate["paragraph_index"]
        if paragraph_index in applied_paragraph_indices:
            candidate_results.append(
                _candidate_result(candidate, "SKIPPED_CHAPTER_CONFLICT", False, "同一段落已有章節候選套用")
            )
            continue
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            candidate_results.append(
                _candidate_result(candidate, "CHAPTER_PARAGRAPH_INDEX_INVALID", False, "章節候選段落索引不存在")
            )
            continue

        paragraph = paragraphs[paragraph_index]
        if not paragraph.text.strip():
            candidate_results.append(
                _candidate_result(candidate, "CHAPTER_PARAGRAPH_EMPTY", False, "章節候選目標段落為空白")
            )
            continue

        try:
            paragraph.style = HEADING_STYLE_NAME
        except Exception as exc:
            candidate_results.append(
                _candidate_result(candidate, "CHAPTER_STYLE_APPLY_FAILED", False, f"章節樣式套用失敗：{exc}")
            )
            continue

        applied_paragraph_indices.add(paragraph_index)
        candidate_results.append(
            _candidate_result(candidate, "APPLIED_CHAPTER_HEADING", True, "章節候選已套用 Heading 2")
        )


def _apply_paragraph_merge_candidates(
    document,
    raw_paragraph_merge_candidates: list[Any],
    candidate_results: list[ReviewApplyCandidateResult],
) -> None:
    candidates_with_order: list[tuple[int, dict[str, Any]]] = []
    for order, item in enumerate(raw_paragraph_merge_candidates):
        candidate, result = _validate_paragraph_merge_candidate(item)
        if result is not None:
            candidate_results.append(result)
            continue
        assert candidate is not None
        candidates_with_order.append((order, candidate))

    used_paragraph_indices: set[int] = set()
    for _order, candidate in sorted(
        candidates_with_order,
        key=lambda pair: (pair[1]["paragraph_index"], -pair[0]),
        reverse=True,
    ):
        paragraph_index = candidate["paragraph_index"]
        next_paragraph_index = candidate["next_paragraph_index"]
        if paragraph_index in used_paragraph_indices or next_paragraph_index in used_paragraph_indices:
            candidate_results.append(
                _candidate_result(candidate, "PARAGRAPH_MERGE_CONFLICT", False, "段落已被其他 merge candidate 使用。")
            )
            continue

        paragraphs = document.paragraphs
        if paragraph_index < 0 or next_paragraph_index < 0 or next_paragraph_index >= len(paragraphs):
            candidate_results.append(
                _candidate_result(candidate, "PARAGRAPH_MERGE_PARAGRAPH_INDEX_INVALID", False, "merge candidate 段落索引不存在。")
            )
            continue
        if next_paragraph_index != paragraph_index + 1:
            candidate_results.append(
                _candidate_result(candidate, "PARAGRAPH_MERGE_NOT_ADJACENT", False, "merge candidate 只能合併相鄰兩段。")
            )
            continue

        paragraph = paragraphs[paragraph_index]
        next_paragraph = paragraphs[next_paragraph_index]
        previous_text = paragraph.text
        next_text = next_paragraph.text
        if not previous_text.strip() or not next_text.strip():
            candidate_results.append(
                _candidate_result(candidate, "PARAGRAPH_MERGE_PARAGRAPH_EMPTY", False, "merge candidate 目標段落不可為空白。")
            )
            continue
        if _is_heading_paragraph(paragraph) or _is_heading_paragraph(next_paragraph):
            candidate_results.append(
                _candidate_result(candidate, "PARAGRAPH_MERGE_CHAPTER_BOUNDARY", False, "merge candidate 不可跨章節標題段落。")
            )
            continue
        if previous_text != candidate["source_text"] or next_text != candidate["next_source_text"]:
            detail_codes = []
            if previous_text != candidate["source_text"]:
                detail_codes.append(PARAGRAPH_MERGE_SOURCE_TEXT_MISMATCH)
            if next_text != candidate["next_source_text"]:
                detail_codes.append(PARAGRAPH_MERGE_NEXT_SOURCE_TEXT_MISMATCH)
            candidate_results.append(
                _candidate_result(
                    candidate,
                    "PARAGRAPH_MERGE_SOURCE_MISMATCH",
                    False,
                    "merge candidate 來源文字與目前文件不符。",
                    detail_code="|".join(detail_codes),
                )
            )
            continue

        try:
            paragraph.text = previous_text.rstrip() + next_text.lstrip()
            _delete_paragraph(next_paragraph)
        except Exception as exc:
            candidate_results.append(
                _candidate_result(candidate, "PARAGRAPH_MERGE_APPLY_FAILED", False, f"段落合併失敗：{exc}")
            )
            continue

        used_paragraph_indices.add(paragraph_index)
        used_paragraph_indices.add(next_paragraph_index)
        candidate_results.append(_candidate_result(candidate, "APPLIED_PARAGRAPH_MERGE", True, "段落合併已套用。"))


def _validate_chapter_candidate(item: Any) -> tuple[dict[str, Any] | None, ReviewApplyCandidateResult | None]:
    if not isinstance(item, dict):
        return None, ReviewApplyCandidateResult("", "", "SKIPPED_INVALID_CANDIDATE", None, False, "chapter candidate 格式不合法")
    candidate_id = item.get("candidate_id")
    candidate_type = item.get("type")
    status = item.get("status")
    paragraph_index = item.get("paragraph_index")
    minimal = {
        "candidate_id": candidate_id,
        "status": status,
        "paragraph_index": paragraph_index if isinstance(paragraph_index, int) else None,
    }
    if not isinstance(candidate_id, str) or not candidate_id.strip() or not isinstance(status, str) or not isinstance(paragraph_index, int):
        return None, _candidate_result(minimal, "SKIPPED_INVALID_CANDIDATE", False, "chapter candidate 必要欄位不合法")
    candidate = {
        "candidate_id": candidate_id,
        "type": candidate_type,
        "status": status,
        "paragraph_index": paragraph_index,
    }
    if candidate_type != "chapter":
        return None, _candidate_result(candidate, "SKIPPED_UNSUPPORTED_TYPE", False, "不支援的 chapter candidate type")
    if status != "accepted":
        return None, _candidate_result(candidate, "SKIPPED_CHAPTER_STATUS", False, "chapter candidate status 不套用")
    return candidate, None


def _validate_paragraph_merge_candidate(item: Any) -> tuple[dict[str, Any] | None, ReviewApplyCandidateResult | None]:
    if not isinstance(item, dict):
        return None, ReviewApplyCandidateResult("", "", "SKIPPED_INVALID_CANDIDATE", None, False, "paragraph merge candidate 格式不合法")

    candidate_id = item.get("candidate_id")
    candidate_type = item.get("type")
    status = item.get("status")
    paragraph_index = item.get("paragraph_index")
    next_paragraph_index = item.get("next_paragraph_index")
    source_text = item.get("source_text")
    next_source_text = item.get("next_source_text")
    minimal = {
        "candidate_id": candidate_id,
        "status": status,
        "paragraph_index": paragraph_index if isinstance(paragraph_index, int) else None,
    }
    if (
        not isinstance(candidate_id, str)
        or not candidate_id.strip()
        or not isinstance(candidate_type, str)
        or not isinstance(status, str)
        or not isinstance(paragraph_index, int)
        or not isinstance(next_paragraph_index, int)
        or not isinstance(source_text, str)
        or not isinstance(next_source_text, str)
    ):
        return None, _candidate_result(minimal, "SKIPPED_INVALID_CANDIDATE", False, "paragraph merge candidate 必要欄位不合法。")

    candidate = {
        "candidate_id": candidate_id,
        "type": candidate_type,
        "status": status,
        "paragraph_index": paragraph_index,
        "next_paragraph_index": next_paragraph_index,
        "source_text": source_text,
        "next_source_text": next_source_text,
    }
    if candidate_type != "paragraph_merge":
        return None, _candidate_result(candidate, "SKIPPED_UNSUPPORTED_TYPE", False, "不支援的 paragraph merge candidate type。")
    if status != "accepted":
        return None, _candidate_result(candidate, "SKIPPED_PARAGRAPH_MERGE_STATUS", False, "paragraph merge candidate status 不套用。")
    return candidate, None


def _validate_candidate(item: Any) -> tuple[dict[str, Any] | None, ReviewApplyCandidateResult | None]:
    if not isinstance(item, dict):
        return None, ReviewApplyCandidateResult("", "", "SKIPPED_INVALID_CANDIDATE", None, False, "candidate 必須是物件。")
    candidate_id = item.get("candidate_id")
    candidate_type = item.get("type")
    status = item.get("status")
    source_text = item.get("source_text")
    resolved_text = item.get("resolved_text")
    paragraph_index = item.get("paragraph_index")
    char_start = item.get("char_start")
    char_end = item.get("char_end")
    context_before = item.get("context_before", "")
    context_after = item.get("context_after", "")
    minimal = {
        "candidate_id": candidate_id,
        "status": status,
        "paragraph_index": paragraph_index if isinstance(paragraph_index, int) else None,
    }
    if (
        not isinstance(candidate_id, str)
        or not candidate_id.strip()
        or not isinstance(source_text, str)
        or not source_text
        or not isinstance(status, str)
        or not isinstance(resolved_text, str)
        or not isinstance(paragraph_index, int)
        or not isinstance(char_start, int)
        or not isinstance(char_end, int)
        or not isinstance(context_before, str)
        or not isinstance(context_after, str)
    ):
        return None, _candidate_result(minimal, "SKIPPED_INVALID_CANDIDATE", False, "candidate 欄位缺失或型別錯誤。")
    candidate = {
        "candidate_id": candidate_id,
        "type": candidate_type,
        "status": status,
        "source_text": source_text,
        "resolved_text": resolved_text,
        "paragraph_index": paragraph_index,
        "char_start": char_start,
        "char_end": char_end,
        "context_before": context_before,
        "context_after": context_after,
    }
    if candidate_type != "high_risk_term":
        return None, _candidate_result(candidate, "SKIPPED_UNSUPPORTED_TYPE", False, "不支援的 candidate type。")
    if status != "accepted":
        return None, _candidate_result(candidate, "SKIPPED_STATUS", False, "status 不允許套用。")
    if not resolved_text.strip():
        return None, _candidate_result(candidate, "SKIPPED_EMPTY_RESOLVED_TEXT", False, "resolved_text 不可空白。")
    return candidate, None


def _candidate_result(
    candidate: dict[str, Any],
    code: str,
    applied: bool,
    message: str,
    detail_code: str = "",
) -> ReviewApplyCandidateResult:
    paragraph_index = candidate.get("paragraph_index")
    return ReviewApplyCandidateResult(
        candidate_id=str(candidate.get("candidate_id", "")),
        status=str(candidate.get("status", "")),
        result_code=code,
        paragraph_index=paragraph_index if isinstance(paragraph_index, int) else None,
        applied=applied,
        message=message,
        detail_code=detail_code,
    )


def _has_conflict(existing: list[tuple[int, int]], char_start: int, char_end: int) -> bool:
    return any(char_start < used_end and char_end > used_start for used_start, used_end in existing)


def _is_skipped_result(result_code: str) -> bool:
    return result_code.startswith("SKIPPED") or result_code in {
        "PARAGRAPH_MERGE_CONFLICT",
        "PARAGRAPH_MERGE_CHAPTER_BOUNDARY",
    }


def _is_heading_paragraph(paragraph) -> bool:
    return getattr(paragraph.style, "name", "") == HEADING_STYLE_NAME


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def _next_reviewed_output_path(output_dir: Path, stem: str) -> Path:
    base = output_dir / f"{stem}_reviewed.docx"
    if not base.exists():
        return base
    index = 1
    while True:
        candidate = output_dir / f"{stem}_reviewed_{index:03d}.docx"
        if not candidate.exists():
            return candidate
        index += 1


def _reviewed_output_path(
    reviewed_output_path: Path | None,
    output_dir: Path | None,
    stem: str,
) -> Path | ReviewApplyResult:
    if reviewed_output_path is not None:
        path = reviewed_output_path.resolve()
        if path.is_dir():
            return _flow_error("REVIEWED_OUTPUT_IS_DIRECTORY", f"reviewed 輸出路徑是資料夾：{path}", str(path))
        if path.suffix.lower() != ".docx":
            return _flow_error("REVIEWED_OUTPUT_NOT_DOCX", f"reviewed 輸出路徑必須是 .docx：{path}", str(path))
        if path.exists():
            return _flow_error("REVIEWED_OUTPUT_EXISTS", f"reviewed 輸出檔已存在：{path}", str(path))
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
        except Exception as exc:
            return _flow_error("REVIEWED_OUTPUT_DIR_CREATE_FAILED", f"無法建立 reviewed 輸出資料夾：{path.parent}", str(exc))
        return path

    if output_dir is None:
        return _flow_error("OUTPUT_DIR_REQUIRED", "未提供輸出資料夾。", "")

    try:
        output_dir.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        return _flow_error("OUTPUT_DIR_CREATE_FAILED", f"無法建立輸出資料夾：{output_dir}", str(exc))
    return _next_reviewed_output_path(output_dir, stem)


def _flow_error(code: str, message: str, technical_detail: str) -> ReviewApplyResult:
    return _flow_error_record(ErrorRecord(code=code, message=message, technical_detail=technical_detail))


def _flow_error_record(error: ErrorRecord) -> ReviewApplyResult:
    return ReviewApplyResult(
        schema=ReviewSchema(errors=[error]),
        output_path=None,
        applied=False,
        candidate_results=[],
        applied_count=0,
        skipped_count=0,
        failed_count=0,
        paragraph_merge_summary=ParagraphMergeSummary(),
        paragraph_merge_diagnostics=ParagraphMergeDiagnostics(),
    )
