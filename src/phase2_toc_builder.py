from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.frozen_spec_v1 import HEADING_STYLE_NAME
from src.review_schema import TocState


TOC_STATUS_NOT_REQUESTED = "not_requested"
TOC_STATUS_SKIPPED_NO_HEADINGS = "skipped_no_headings"
TOC_STATUS_SKIPPED_EXISTING_TOC = "skipped_existing_toc"
TOC_STATUS_FIELD_INSERTED = "field_inserted"
TOC_STATUS_FALLBACK_CHAPTER_LIST = "fallback_chapter_list"
TOC_STATUS_FAILED = "failed"


def insert_minimal_toc(document, requested: bool) -> TocState:
    if not requested:
        return TocState(requested=False, status=TOC_STATUS_NOT_REQUESTED, fallback_used=False, chapter_count=0)

    headings = _heading2_texts(document)
    if _has_existing_toc_title(document):
        return TocState(
            requested=True,
            status=TOC_STATUS_SKIPPED_EXISTING_TOC,
            fallback_used=False,
            chapter_count=len(headings),
        )
    if not headings:
        return TocState(
            requested=True,
            status=TOC_STATUS_SKIPPED_NO_HEADINGS,
            fallback_used=False,
            chapter_count=0,
        )

    try:
        _insert_toc_field(document)
        return TocState(
            requested=True,
            status=TOC_STATUS_FIELD_INSERTED,
            fallback_used=False,
            chapter_count=len(headings),
        )
    except Exception:
        try:
            _insert_fallback_chapter_list(document, headings)
        except Exception:
            return TocState(
                requested=True,
                status=TOC_STATUS_FAILED,
                fallback_used=True,
                chapter_count=len(headings),
            )
        return TocState(
            requested=True,
            status=TOC_STATUS_FALLBACK_CHAPTER_LIST,
            fallback_used=True,
            chapter_count=len(headings),
        )


def _heading2_texts(document) -> list[str]:
    return [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.strip() and _style_name(paragraph) == HEADING_STYLE_NAME
    ]


def _has_existing_toc_title(document) -> bool:
    for paragraph in document.paragraphs[:5]:
        text = paragraph.text.strip()
        if text:
            return text == "目錄"
    return False


def _insert_toc_field(document) -> None:
    title = document.add_paragraph("目錄")
    toc_paragraph = document.add_paragraph()
    run = toc_paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "2-2" \h \z \u'
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    toc_paragraph.add_run("請在 Word 中更新目錄。")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    toc_paragraph.runs[-1]._r.append(end)

    _move_to_document_start(document, [title._p, toc_paragraph._p])


def _insert_fallback_chapter_list(document, headings: list[str]) -> None:
    paragraphs = [document.add_paragraph("目錄")]
    paragraphs.extend(document.add_paragraph(text) for text in headings)
    _move_to_document_start(document, [paragraph._p for paragraph in paragraphs])


def _move_to_document_start(document, paragraph_elements) -> None:
    body = document._body._element
    first = body[0] if len(body) else None
    for element in paragraph_elements:
        body.remove(element)
        if first is None:
            body.append(element)
        else:
            body.insert(body.index(first), element)


def _style_name(paragraph) -> str:
    style = paragraph.style
    return style.name if style is not None else ""
